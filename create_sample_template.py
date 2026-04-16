from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_sample():
    doc = Document()
    
    # Page margins
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.0)

    # Header
    p1 = doc.add_paragraph('BỘ CÔNG THƯƠNG\nTRƯỜNG ĐẠI HỌC ĐIỆN LỰC')
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p1.runs:
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'

    doc.add_paragraph()
    
    # Title
    t = doc.add_paragraph('ĐỀ CƯƠNG CHI TIẾT HỌC PHẦN')
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run('\n[TEN_VIET]')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()

    # Section 1
    h1 = doc.add_paragraph('1. Thông tin chung về học phần')
    h1.runs[0].bold = True
    
    doc.add_paragraph('- Mã học phần: [MA_HP]')
    doc.add_paragraph('- Số tín chỉ: [SO_TIN_CHI]')
    doc.add_paragraph('- Trình độ đào tạo: [TRINH_DO]')
    doc.add_paragraph('- Loại học phần: [LOAI_HP]')
    
    # Table example for Teachers
    doc.add_paragraph('Giảng viên phụ trách chính:')
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'STT'
    hdr_cells[1].text = 'Họ và tên'
    hdr_cells[2].text = 'Điện thoại'
    hdr_cells[3].text = 'Email'
    
    # Template Row
    row = table.rows[1].cells
    row[0].text = '[GV_CHINH_STT]'
    row[1].text = '[GV_CHINH_HO_TEN]'
    row[2].text = '[GV_CHINH_SDT]'
    row[3].text = '[GV_CHINH_EMAIL]'

    doc.add_paragraph()

    # Section 2
    h2 = doc.add_paragraph('2. Mô tả tóm tắt nội dung học phần')
    h2.runs[0].bold = True
    doc.add_paragraph('[MO_TA]')

    doc.add_paragraph()

    # Content Table (Lý thuyết)
    h3 = doc.add_paragraph('3. Nội dung chi tiết học phần (Lý thuyết)')
    h3.runs[0].bold = True
    
    nt = doc.add_table(rows=2, cols=3)
    nt.style = 'Table Grid'
    h = nt.rows[0].cells
    h[0].text = 'Tên Chương / Bài học'
    h[1].text = 'Số tiết LT'
    h[2].text = 'Chuẩn đầu ra'
    
    r = nt.rows[1].cells
    r[0].text = '[ND_LT_TEN]'
    r[1].text = '[ND_LT_GIO_LT]'
    r[2].text = '[ND_LT_CDR_MA]'

    doc.add_paragraph()

    # Signature Block
    sig = doc.add_table(rows=2, cols=2)
    # Right align the first paragraph for date
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_date.add_run('[DIA_DIEM_KY], [NGAY_KY]').italic = True
    
    c1 = sig.rows[0].cells[0]
    c1.text = '[CHUC_DANH_KY_TRAI]'
    c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c1.paragraphs[0].runs[0].bold = True
    
    c2 = sig.rows[0].cells[1]
    c2.text = '[CHUC_DANH_KY_PHAI]'
    c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c2.paragraphs[0].runs[0].bold = True
    
    sig.rows[1].cells[0].text = '\n\n\n\n[HO_TEN_KY_TRAI]'
    sig.rows[1].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig.rows[1].cells[0].paragraphs[0].runs[0].bold = True
    
    sig.rows[1].cells[1].text = '\n\n\n\n[HO_TEN_KY_PHAI]'
    sig.rows[1].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig.rows[1].cells[1].paragraphs[0].runs[0].bold = True

    # Save
    out = os.path.join("g:\\My Drive\\01. Working\\EPU\\EPU Utis\\QLCTDT Anti\\V1.1\\templates", "Mau_De_Cuong_Chuan.docx")
    doc.save(out)
    print(f"Created sample template at: {out}")

if __name__ == '__main__':
    create_sample()
