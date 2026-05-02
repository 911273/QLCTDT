# templates/styles.py
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

FONT_NAME = "Times New Roman"
FONT_SIZE_NORMAL = Pt(12)
FONT_SIZE_TITLE = Pt(14)
COLOR_HEADER_BG = RGBColor(0xD9, 0xD9, 0xD9)
COLOR_ERROR_BG  = RGBColor(0xFF, 0xCC, 0xCC)  # đỏ nhạt cho lỗi
COLOR_WARN_BG   = RGBColor(0xFF, 0xFF, 0xCC)  # vàng nhạt cho warning

def apply_cell_style(cell, bold=False, center=False, bg_color=None, font_size=None, font_color=None):
    """Áp dụng style cho cell của bảng (Background, Center, Bold)."""
    # Xóa text rác nếu có (dù table.add_row() thường tạo cell trống, nhưng nó tạo 1 paragraph trống)
    if not cell.paragraphs:
        cell.add_paragraph()
    
    for p in cell.paragraphs:
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            if bold:
                run.font.bold = True
            if font_size:
                run.font.size = font_size
            if font_color:
                run.font.color.rgb = font_color
                
    # Background color
    if bg_color:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        # bg_color is an RGBColor object or hex string. E.g RGBColor(0xD9, 0xD9, 0xD9) -> D9D9D9
        if isinstance(bg_color, RGBColor):
            hex_color = '%02X%02X%02X' % (bg_color[0], bg_color[1], bg_color[2])
        else:
            hex_color = str(bg_color).replace('#', '')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

def set_table_style(table, autofit=True, repeat_header=True):
    """
    Áp dụng style cho bảng:
    - AutoFit to window
    - Repeat Header Rows
    - Table Grid (Viền solid 0.5pt đen)
    """
    table.style = 'Table Grid'
    
    if autofit:
        table.autofit = True
        # Force docx to make table 100% width
        from docx.oxml.shared import OxmlElement
        from docx.oxml.ns import qn
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:type'), 'pct')
        tblW.set(qn('w:w'), '5000') # 5000 pct = 100% in OOXML
        tblPr.append(tblW)
        
    if repeat_header:
        set_repeat_header(table)

def set_repeat_header(table):
    """Bật Repeat Header Rows cho list các header (thường là row 0)."""
    if len(table.rows) > 0:
        tr = table.rows[0]._tr
        trPr = tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:tblHeader')
        tblHeader.set(qn('w:val'), '1')
        trPr.append(tblHeader)

def enforce_font_on_doc(doc):
    """Đảm bảo quét qua toàn bộ document, ép tất cả về Times New Roman."""
    for p in doc.paragraphs:
        for run in p.runs:
            run.font.name = FONT_NAME
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    for run in p.runs:
                        run.font.name = FONT_NAME
