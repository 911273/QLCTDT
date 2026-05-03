# statistics_dialog.py
import tkinter as tk
from tkinter import filedialog
import threading
import re
from datetime import datetime

from utils.ui_utils import (show_modern_info, show_modern_warning, 
                             show_modern_error, ask_modern_yesno)

# Matplotlib for analytics
try:
    import matplotlib.pyplot as plt
    from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
    HAS_MATPLOTLIB = True
except ImportError:
    HAS_MATPLOTLIB = False

# Excel export
try:
    import openpyxl
    from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

# Word export
try:
    import docx
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

import ttkbootstrap as tb
from ttkbootstrap.constants import *
from ttkbootstrap.widgets import Meter

# Màu sắc cho I/R/M
IRM_COLORS = {
    'I': '#FFF2CC', # Vàng nhạt
    'R': '#FCE4D6', # Cam nhạt
    'M': '#E2EFDA'  # Xanh lá nhạt
}

IRM_ICONS = {
    'I': '🟨 I', 
    'R': '🟧 R', 
    'M': '🟩 M', 
    '': '-'
}

class StatisticsDialog(tb.Toplevel):
    def __init__(self, parent, db):
        super().__init__(title="📊 Báo cáo Kiểm định chất lượng CTĐT", size=(1300, 900))
        self.db = db
        self.parent = parent
        
        self.matrix_data = []  # Toàn bộ CLO của CTĐT
        self.plo_list = []     # Danh sách mã PLO
        self.irm_stats = {}    # PLO -> {'I': x, 'R': y, 'M': z}
        self.bloom_counts = [0] * 6
        self.ctdt_info = {}    # Thông tin CTĐT hiện tại
        
        self._build_ui()
        self.position_center()
        self.grab_set()
        
        # Load danh sách CTĐT vào combobox ngay khi mở
        self._refresh_ctdt_combobox()

    def _build_ui(self):
        # Header area
        header = tb.Frame(self, padding=10)
        header.pack(fill=X)
        
        tb.Label(header, text="Chương trình đào tạo:", font=('Arial', 10, 'bold')).pack(side=LEFT, padx=5)
        self.cbo_ctdt = tb.Combobox(header, state="readonly", width=40)
        self.cbo_ctdt.pack(side=LEFT, padx=5)
        
        btn_load = tb.Button(header, text="🔍 Chạy Phân tích", bootstyle=SUCCESS, command=self.refresh_all)
        btn_load.pack(side=LEFT, padx=10)
        
        # Main Notebook
        self.nb = tb.Notebook(self, bootstyle=PRIMARY)
        self.nb.pack(fill=BOTH, expand=YES, padx=10, pady=5)
        
        # Tabs
        self.tab_matrix = self._init_tab_matrix()
        self.nb.add(self.tab_matrix, text="📐 Ma trận CLO-PLO")
        
        self.tab_bloom = self._init_tab_bloom()
        self.nb.add(self.tab_bloom, text="🌸 Bao phủ mức Bloom")
        
        self.tab_irm = self._init_tab_irm_plo()
        self.nb.add(self.tab_irm, text="🔄 Phân bổ I/R/M theo PLO")

        # Bottom Bar
        footer = tb.Frame(self, padding=10)
        footer.pack(fill=X)
        
        self.lbl_status = tb.Label(footer, text="Sẵn sàng", foreground='grey')
        self.lbl_status.pack(side=LEFT)
        
        tb.Button(footer, text="❌ Đóng", bootstyle=SECONDARY, command=self.destroy).pack(side=RIGHT, padx=5)
        tb.Button(footer, text="📑 Xuất Báo cáo Kiểm định (Word)", bootstyle=PRIMARY, 
                  command=self._export_word_report).pack(side=RIGHT, padx=5)
        tb.Button(footer, text="📊 Xuất Ma trận (Excel)", bootstyle=INFO, 
                  command=self._export_matrix_excel).pack(side=RIGHT, padx=5)

    def _init_tab_matrix(self):
        tab = tb.Frame(self.nb, padding=10)
        
        # 1. Nhãn tổng kết (Dưới cùng)
        self.lbl_matrix_summary = tb.Label(tab, text="Tổng số: 0 CLO | Chưa phân tích", 
                                           font=('Arial', 10, 'italic'), foreground='blue')
        self.lbl_matrix_summary.pack(side=BOTTOM, anchor='w', pady=5)
        
        # 2. Scrollbar ngang (Trên nhãn tổng kết)
        hsb = tb.Scrollbar(tab, orient=HORIZONTAL)
        hsb.pack(side=BOTTOM, fill=X)
        
        # 3. Container cho Treeview và Scrollbar dọc
        table_frame = tb.Frame(tab)
        table_frame.pack(fill=BOTH, expand=YES)
        
        self.tree_matrix = tb.Treeview(table_frame, show='headings', bootstyle=INFO)
        vsb = tb.Scrollbar(table_frame, orient=VERTICAL, command=self.tree_matrix.yview)
        
        # Cấu hình liên kết scrollbar
        self.tree_matrix.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        hsb.configure(command=self.tree_matrix.xview)
        
        self.tree_matrix.pack(side=LEFT, fill=BOTH, expand=YES)
        vsb.pack(side=RIGHT, fill=Y)
        
        return tab

    def _init_tab_bloom(self):
        tab = tb.Frame(self.nb, padding=10)
        self.bloom_canvas_frame = tb.Frame(tab)
        self.bloom_canvas_frame.pack(fill=BOTH, expand=YES)
        
        if not HAS_MATPLOTLIB:
            tb.Label(self.bloom_canvas_frame, text="Vui lòng cài đặt matplotlib để xem biểu đồ", 
                     foreground='red').pack(pady=50)
        
        self.lbl_bloom_warn = tb.Label(tab, text="", font=('Arial', 10, 'bold'))
        self.lbl_bloom_warn.pack(pady=10)
        
        return tab

    def _init_tab_irm_plo(self):
        tab = tb.Frame(self.nb, padding=10)
        self.irm_canvas_frame = tb.Frame(tab)
        self.irm_canvas_frame.pack(fill=BOTH, expand=YES)
        
        if not HAS_MATPLOTLIB:
            tb.Label(self.irm_canvas_frame, text="Vui lòng cài đặt matplotlib để xem biểu đồ", 
                     foreground='red').pack(pady=50)
            
        self.lbl_irm_warn = tb.Label(tab, text="", font=('Arial', 10, 'bold'))
        self.lbl_irm_warn.pack(pady=10)
        
        return tab

    def _refresh_ctdt_combobox(self):
        try:
            # Lấy danh sách CTĐT từ DB (Bảng chuẩn chuong_trinh_dao_tao)
            rows = self.db.conn.execute("SELECT id, ten, bac FROM chuong_trinh_dao_tao ORDER BY bac, ten").fetchall()
            self.ctdt_map = {f"{r['ten']} ({r['bac']})": r['id'] for r in rows}
            names = list(self.ctdt_map.keys())
            self.cbo_ctdt.config(values=names)
            if names: self.cbo_ctdt.current(0)
        except Exception as e:
            print(f"Error loading CTDT: {e}")

    # -------------------------------------------------------------------------
    # LOGIC: FETCH & PROCESS DATA
    # -------------------------------------------------------------------------
    def refresh_all(self):
        ctdt_name = self.cbo_ctdt.get()
        if not ctdt_name:
            show_modern_warning(self, "Thông báo", "Vui lòng chọn một Chương trình đào tạo.")
            return
            
        self.lbl_status.config(text="Đang xử lý dữ liệu...", foreground='blue')
        threading.Thread(target=self._load_data_async, args=(self.ctdt_map[ctdt_name],), daemon=True).start()

    def _load_data_async(self, ctdt_id):
        try:
            # 1. Lấy thông tin CTĐT
            ctdt_row = self.db.conn.execute("SELECT * FROM chuong_trinh_dao_tao WHERE id = ?", (ctdt_id,)).fetchone()
            self.ctdt_info = dict(ctdt_row) if ctdt_row else {}

            # 2. Lấy danh sách PLO
            plos = self.db.conn.execute("SELECT ma, mo_ta FROM ctdt_plo WHERE ctdt_id = ? ORDER BY ma", (ctdt_id,)).fetchall()
            self.plo_list = [p['ma'] for p in plos]
            if not self.plo_list:
                # Fallback nếu chưa định nghĩa PLO
                self.plo_list = [f"PLO{i}" for i in range(1, 13)]

            # 3. Lấy dữ liệu học phần và CLO
            hps = self.db.conn.execute("""
                SELECT hp.id, hp.ma, hp.ten_viet 
                FROM hoc_phan hp 
                JOIN ctdt_hoc_phan ch ON hp.id = ch.hp_id
                WHERE ch.ctdt_id = ?
                ORDER BY ch.thu_tu, hp.ma
            """, (ctdt_id,)).fetchall()

            self.matrix_data = []
            self.irm_stats = {plo: {'I': 0, 'R': 0, 'M': 0, 'total': 0} for plo in self.plo_list}
            self.bloom_counts = [0] * 6
            
            from services.deccuong_service import DeCuongValidator
            validator = DeCuongValidator(self.db)

            for hp in hps:
                hp_id, hp_ma, hp_ten = hp['id'], hp['ma'], hp['ten_viet']
                clos_rows = self.db.conn.execute("""
                    SELECT ma, mo_ta, cdr_ma, level_irm 
                    FROM clo 
                    WHERE hp_id = ? AND (la_tieu_de_nhom IS NULL OR la_tieu_de_nhom = 0)
                    ORDER BY id
                """, (hp_id,)).fetchall()

                for clo in clos_rows:
                    clo_ma, clo_mota, cdr_ma, level = clo['ma'], clo['mo_ta'], clo['cdr_ma'], clo['level_irm']
                    if not clo_ma: continue
                    
                    # A. Nhận diện mức Bloom
                    desc = (clo_mota or "").lower().strip()
                    bloom_level = 1 # Mặc định Nhớ
                    for lvl, verbs in validator.DONG_TU_BLOOM.items():
                        if any(desc.startswith(v) for v in verbs):
                            bloom_level = lvl
                            break
                    self.bloom_counts[bloom_level-1] += 1

                    # B. Phân tích PLO mapping
                    # Hỗ trợ nhiều PLO cách nhau bởi dấu phẩy
                    mapped_plos = [p.strip().upper() for p in re.split(r'[,\s;]+', str(cdr_ma)) if p.strip()]
                    lvl_val = str(level).strip().upper() if level else 'I'
                    if lvl_val not in ['I', 'R', 'M']: lvl_val = 'I'
                    
                    row_dict = {
                        'hp_ma': hp_ma,
                        'hp_ten': hp_ten,
                        'clo_ma': clo_ma,
                        'clo_mota': clo_mota,
                        'bloom': bloom_level,
                        'plos': {}
                    }

                    for p_code in mapped_plos:
                        if p_code in self.plo_list:
                            row_dict['plos'][p_code] = lvl_val
                            self.irm_stats[p_code][lvl_val] += 1
                            self.irm_stats[p_code]['total'] += 1
                            
                    self.matrix_data.append(row_dict)

            # Cập nhật UI
            self.after(0, self._render_all_tabs)
            
        except Exception as e:
            self.after(0, lambda: self.lbl_status.config(text=f"Lỗi: {str(e)}", foreground='red'))
            print(f"Stats Processing Error: {e}")

    def _render_all_tabs(self):
        self._render_matrix_ui()
        self._render_bloom_chart()
        self._render_irm_chart()
        self.lbl_status.config(text=f"Đã tải dữ liệu lúc {datetime.now().strftime('%H:%M:%S')}", foreground='green')

    def _render_matrix_ui(self):
        # Thiết lập cột
        cols = ["hp", "clo", "bloom"] + self.plo_list
        self.tree_matrix["columns"] = cols
        
        self.tree_matrix.heading("hp", text="Học phần")
        self.tree_matrix.heading("clo", text="Mã CLO")
        self.tree_matrix.heading("bloom", text="Bloom")
        
        self.tree_matrix.column("hp", width=220, anchor=W)
        self.tree_matrix.column("clo", width=80, anchor=CENTER)
        self.tree_matrix.column("bloom", width=60, anchor=CENTER)
        
        for p in self.plo_list:
            # Rút gọn header nếu quá dài
            disp_p = p if len(p) < 8 else p[:6] + ".."
            self.tree_matrix.heading(p, text=disp_p)
            self.tree_matrix.column(p, width=55, anchor=CENTER)
            
        # Xóa dữ liệu cũ
        self.tree_matrix.delete(*self.tree_matrix.get_children())
        
        current_hp = None
        for data in self.matrix_data:
            hp_disp = f"{data['hp_ma']} - {data['hp_ten']}" if data['hp_ma'] != current_hp else ""
            current_hp = data['hp_ma']
            
            row = [hp_disp, data['clo_ma'], f"L{data['bloom']}"]
            for p in self.plo_list:
                val = data['plos'].get(p, '')
                row.append(IRM_ICONS.get(val, val))
                
            self.tree_matrix.insert("", END, values=row)

        # Thêm dòng tổng kết (Số CLO mỗi PLO)
        summary_row = ["TỔNG SỐ CLO", "", ""]
        for p in self.plo_list:
            summary_row.append(str(self.irm_stats[p]['total']))
        self.tree_matrix.insert("", END, values=summary_row, tags=('summary',))
        self.tree_matrix.tag_configure('summary', background='#f0f0f0', font=('Arial', 9, 'bold'))

        self.lbl_matrix_summary.config(text=f"📊 Tổng số: {len(self.matrix_data)} CLO từ {len(set(d['hp_ma'] for d in self.matrix_data))} học phần.")

    def _render_bloom_chart(self):
        if not HAS_MATPLOTLIB: return
        for child in self.bloom_canvas_frame.winfo_children(): child.destroy()
        
        fig, ax = plt.subplots(figsize=(8, 4), dpi=100)
        levels = ["L1", "L2", "L3", "L4", "L5", "L6"]
        full_names = ["Nhớ", "Hiểu", "Vận dụng", "Phân tích", "Đánh giá", "Sáng tạo"]
        colors = ['#ff9999','#ffcc99','#66b3ff','#99ff99','#ffb3e6','#c2c2f0']
        
        bars = ax.bar(levels, self.bloom_counts, color=colors, edgecolor='grey', alpha=0.8)
        ax.set_title("Phân bổ mức độ nhận thức (Bloom) toàn CTĐT", fontsize=12, pad=15)
        ax.set_ylabel("Số lượng CLO")
        
        # Thêm text chú thích cho các mức
        for i, bar in enumerate(bars):
            yval = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2, yval + 0.1, f"{int(yval)}", ha='center', va='bottom', fontweight='bold')
            ax.text(bar.get_x() + bar.get_width()/2, -0.5, full_names[i], ha='center', va='top', fontsize=8)

        fig.subplots_adjust(bottom=0.2)
        # fig.tight_layout() # Bỏ tight_layout vì gây cảnh báo decor decorations
        canvas = FigureCanvasTkAgg(fig, master=self.bloom_canvas_frame)
        canvas.draw()
        canvas.get_tk_widget().pack(fill=BOTH, expand=YES)
        
        # Cảnh báo nếu thiếu mức
        missing = [f"L{i+1}" for i, v in enumerate(self.bloom_counts) if v == 0]
        if missing:
            self.lbl_bloom_warn.config(text=f"⚠️ Cảnh báo: CTĐT đang thiếu các mức Bloom: {', '.join(missing)}", foreground='red')
        else:
            self.lbl_bloom_warn.config(text="✅ Tuyệt vời! CTĐT bao phủ đầy đủ cả 6 mức độ nhận thức.", foreground='#28a745')

    def _render_irm_chart(self):
        if not HAS_MATPLOTLIB: return
        for child in self.irm_canvas_frame.winfo_children(): child.destroy()
        
        if not self.plo_list: return

        plos = self.plo_list
        i_vals = [self.irm_stats[p]['I'] for p in plos]
        r_vals = [self.irm_stats[p]['R'] for p in plos]
        m_vals = [self.irm_stats[p]['M'] for p in plos]

        fig, ax = plt.subplots(figsize=(10, 4), dpi=100)
        
        # Vẽ stacked bar
        ax.bar(plos, i_vals, label='I (Giới thiệu)', color='#f1c40f', alpha=0.8)
        ax.bar(plos, r_vals, bottom=i_vals, label='R (Củng cố)', color='#e67e22', alpha=0.8)
        bottom_m = [i+r for i,r in zip(i_vals, r_vals)]
        ax.bar(plos, m_vals, bottom=bottom_m, label='M (Thành thạo)', color='#2ecc71', alpha=0.8)

        ax.set_title("Ma trận bao phủ I/R/M theo từng PLO", fontsize=12, pad=15)
        ax.set_ylabel("Số lượng CLO đóng góp")
        ax.legend(loc='upper right', frameon=True)
        
        plt.xticks(rotation=45, ha='right')
        fig.subplots_adjust(bottom=0.25, top=0.9, left=0.1, right=0.9)
        # fig.tight_layout()
        canvas = FigureCanvasTkAgg(fig, master=self.irm_canvas_frame)
        canvas.draw()
        canvas.get_tk_widget().pack(fill=BOTH, expand=YES)

        # Cảnh báo thiếu M
        missing_m = [p for p in plos if self.irm_stats[p]['M'] == 0]
        if missing_m:
            self.lbl_irm_warn.config(text=f"❌ Cảnh báo đỏ: Các PLO sau thiếu CLO ở mức Thành thạo (M): {', '.join(missing_m)}", foreground='red')
        else:
            self.lbl_irm_warn.config(text="✅ Đạt chuẩn: Tất cả PLO đều được cover ở mức Thành thạo (M).", foreground='#28a745')

    # -------------------------------------------------------------------------
    # EXPORT FUNCTIONS
    # -------------------------------------------------------------------------
    def _export_matrix_excel(self):
        if not HAS_OPENPYXL:
            show_modern_error(self, "Lỗi", "Vui lòng cài đặt openpyxl (pip install openpyxl).")
            return
        
        if not self.matrix_data:
            show_modern_warning(self, "Thông báo", "Vui lòng 'Chạy Phân tích' trước khi xuất.")
            return

        path = filedialog.asksaveasfilename(defaultextension=".xlsx", 
                                            filetypes=[("Excel", "*.xlsx")], 
                                            title="Lưu Ma trận CLO-PLO",
                                            initialfile=f"MaTran_CLO_PLO_{self.ctdt_info.get('id', 'Report')}.xlsx")
        if not path: return
        
        try:
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Ma trận CLO-PLO"
            
            # Styles
            header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
            header_font = Font(color="FFFFFF", bold=True)
            center_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
            border = Border(left=Side(style='thin'), right=Side(style='thin'), 
                           top=Side(style='thin'), bottom=Side(style='thin'))

            # Header row
            cols = ["STT", "Mã Học phần", "Tên Học phần", "Mã CLO", "Mức Bloom"] + self.plo_list
            ws.append(cols)
            for cell in ws[1]:
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = center_align
                cell.border = border

            # Data rows
            fill_i = PatternFill(start_color="FFF2CC", end_color="FFF2CC", fill_type="solid")
            fill_r = PatternFill(start_color="FCE4D6", end_color="FCE4D6", fill_type="solid")
            fill_m = PatternFill(start_color="E2EFDA", end_color="E2EFDA", fill_type="solid")

            for i, data in enumerate(self.matrix_data, 1):
                row_vals = [i, data['hp_ma'], data['hp_ten'], data['clo_ma'], f"L{data['bloom']}"]
                for p in self.plo_list:
                    row_vals.append(data['plos'].get(p, ''))
                ws.append(row_vals)
                
                curr_row = ws.max_row
                for col_idx, p in enumerate(self.plo_list, start=6):
                    cell = ws.cell(row=curr_row, column=col_idx)
                    cell.alignment = center_align
                    cell.border = border
                    val = data['plos'].get(p, '')
                    if val == 'I': cell.fill = fill_i
                    elif val == 'R': cell.fill = fill_r
                    elif val == 'M': cell.fill = fill_m
                
                # Căn lề các cột đầu
                for c in range(1, 6):
                    ws.cell(row=curr_row, column=c).border = border
                    ws.cell(row=curr_row, column=c).alignment = Alignment(vertical="center")

            # Auto-size columns
            ws.column_dimensions['B'].width = 15
            ws.column_dimensions['C'].width = 35
            ws.column_dimensions['D'].width = 10
            
            wb.save(path)
            show_modern_info(self, "Thành công", f"Đã xuất file Excel tại:\n{path}")
        except Exception as e:
            show_modern_error(self, "Lỗi", f"Không thể lưu file Excel: {e}")

    def _export_word_report(self):
        if not HAS_DOCX:
            show_modern_error(self, "Lỗi", "Vui lòng cài đặt python-docx.")
            return
            
        if not self.matrix_data:
            show_modern_warning(self, "Thông báo", "Vui lòng 'Chạy Phân tích' trước.")
            return

        path = filedialog.asksaveasfilename(defaultextension=".docx", 
                                            filetypes=[("Word", "*.docx")], 
                                            title="Xuất Báo cáo Kiểm định",
                                            initialfile=f"BaoCao_KiemDinh_{self.ctdt_info.get('id', 'CTDT')}.docx")
        if not path: return
        
        try:
            doc = docx.Document()
            
            # Tiêu đề Header (Trường/Khoa)
            header_table = doc.add_table(rows=1, cols=2)
            header_table.width = Inches(6.5)
            c1 = header_table.cell(0, 0).paragraphs[0]
            c1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            c1.add_run("TRƯỜNG ĐẠI HỌC ĐIỆN LỰC\nKHOA CÔNG NGHỆ THÔNG TIN").bold = True
            
            c2 = header_table.cell(0, 1).paragraphs[0]
            c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            c2.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\nĐộc lập - Tự do - Hạnh phúc").bold = True
            
            doc.add_paragraph("\n")
            
            # Tiêu đề báo cáo
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title.add_run("BÁO CÁO PHÂN TÍCH CHẤT LƯỢNG CHƯƠNG TRÌNH ĐÀO TẠO")
            run.bold = True
            run.font.size = Pt(16)
            
            info = doc.add_paragraph()
            info.alignment = WD_ALIGN_PARAGRAPH.CENTER
            info.add_run(f"Chương trình: {self.ctdt_info.get('ten', 'N/A')}\nBậc đào tạo: {self.ctdt_info.get('bac', 'N/A')}\nNgày xuất: {datetime.now().strftime('%d/%m/%Y')}")
            
            # 1. Ma trận CLO-PLO
            doc.add_heading('1. Ma trận đóng góp CLO - PLO', level=1)
            doc.add_paragraph("Bảng dưới đây thống kê sự đóng góp của tất cả CLO thuộc các học phần trong CTĐT vào các Chuẩn đầu ra (PLO) theo các mức I (Giới thiệu), R (Củng cố), M (Thành thạo).")
            
            table = doc.add_table(rows=1, cols=len(self.plo_list) + 2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Học phần'
            hdr_cells[1].text = 'CLO'
            for i, p_name in enumerate(self.plo_list):
                hdr_cells[i+2].text = p_name
                
            # Header font size
            for cell in table.rows[0].cells:
                for p in cell.paragraphs:
                    for r in p.runs: r.font.size = Pt(9); r.bold = True

            for data in self.matrix_data:
                row_cells = table.add_row().cells
                row_cells[0].text = data['hp_ma']
                row_cells[1].text = data['clo_ma']
                for i, p_name in enumerate(self.plo_list):
                    val = data['plos'].get(p_name, '')
                    row_cells[i+2].text = val
                    if val:
                        # Tô màu cell (xml)
                        shading_elm = OxmlElement('w:shd')
                        color = IRM_COLORS.get(val, 'FFFFFF').replace('#', '')
                        shading_elm.set(qn('w:fill'), color)
                        row_cells[i+2]._tc.get_or_add_tcPr().append(shading_elm)

            # 2. Thống kê IRM
            doc.add_heading('2. Thống kê bao phủ mức độ Thành thạo (I/R/M)', level=1)
            doc.add_paragraph("Phân tích mức độ hỗ trợ của CTĐT đối với từng Chuẩn đầu ra.")
            
            stat_table = doc.add_table(rows=1, cols=5)
            stat_table.style = 'Table Grid'
            hcells = stat_table.rows[0].cells
            titles = ['PLO', 'Giới thiệu (I)', 'Củng cố (R)', 'Thành thạo (M)', 'Tổng CLO']
            for i, t in enumerate(titles): 
                hcells[i].text = t
                hcells[i].paragraphs[0].runs[0].bold = True
            
            for p_name in self.plo_list:
                rcells = stat_table.add_row().cells
                rcells[0].text = p_name
                rcells[1].text = str(self.irm_stats[p_name]['I'])
                rcells[2].text = str(self.irm_stats[p_name]['R'])
                rcells[3].text = str(self.irm_stats[p_name]['M'])
                rcells[4].text = str(self.irm_stats[p_name]['total'])
                
                if self.irm_stats[p_name]['M'] == 0:
                    rcells[3].paragraphs[0].add_run(" (Thiếu!)").font.color.rgb = RGBColor(255, 0, 0)

            # 3. Bloom
            doc.add_heading('3. Phân bổ mức độ nhận thức (Bloom)', level=1)
            doc.add_paragraph("Thống kê số lượng CLO đạt được các mức độ nhận thức theo thang đo Bloom.")
            
            b_table = doc.add_table(rows=1, cols=2)
            b_table.style = 'Table Grid'
            b_table.rows[0].cells[0].text = 'Mức độ Bloom'
            b_table.rows[0].cells[1].text = 'Số lượng CLO'
            
            levels = ["Nhớ (L1)", "Hiểu (L2)", "Vận dụng (L3)", "Phân tích (L4)", "Đánh giá (L5)", "Sáng tạo (L6)"]
            for i, lvl in enumerate(levels):
                rcells = b_table.add_row().cells
                rcells[0].text = lvl
                rcells[1].text = str(self.bloom_counts[i])
            
            doc.add_paragraph("\nKết luận: " + ("CTĐT đạt yêu cầu về độ bao phủ." if all(v > 0 for v in self.bloom_counts) else "Cần bổ sung các CLO ở mức độ nhận thức cao."))

            # Ký tên
            doc.add_paragraph("\n")
            sign_table = doc.add_table(rows=1, cols=2)
            sign_table.width = Inches(6.5)
            s1 = sign_table.cell(0, 0).paragraphs[0]
            s1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            s1.add_run("Người lập báo cáo\n\n\n(Ký và ghi rõ họ tên)")
            
            s2 = sign_table.cell(0, 1).paragraphs[0]
            s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            s2.add_run("Trưởng bộ môn/Khoa\n\n\n(Ký và ghi rõ họ tên)")

            doc.save(path)
            show_modern_info(self, "Thành công", f"Báo cáo Word đã được lưu:\n{path}")
        except Exception as e:
            show_modern_error(self, "Lỗi", f"Lỗi xuất file Word: {e}")
