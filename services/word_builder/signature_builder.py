# services/word_builder/signature_builder.py
from docx.enum.text import WD_ALIGN_PARAGRAPH
from templates.styles import apply_cell_style, set_table_style, COLOR_HEADER_BG, FONT_SIZE_NORMAL

class SignatureBuilder:
    @staticmethod
    def build(doc, data: dict):
        h8 = doc.add_paragraph()
        run = h8.add_run("8. TIẾN TRÌNH CẬP NHẬT ĐỀ CƯƠNG")
        run.font.bold = True
        run.font.size = FONT_SIZE_NORMAL

        tien_trinh = data.get('tien_trinh', [])
        
        table = doc.add_table(rows=len(tien_trinh) + 1, cols=2)
        set_table_style(table)
        
        table.cell(0, 0).text = "Nội dung cập nhật"
        table.cell(0, 1).text = "Người cập nhật"
        apply_cell_style(table.cell(0, 0), bold=True, center=True, bg_color=COLOR_HEADER_BG)
        apply_cell_style(table.cell(0, 1), bold=True, center=True, bg_color=COLOR_HEADER_BG)

        for i, tt in enumerate(tien_trinh):
            row = table.rows[i + 1]
            row.cells[0].text = f"Lần {tt.get('lan', i+1)}: {tt.get('noi_dung', '')}"
            row.cells[1].text = tt.get('nguoi', '')
            apply_cell_style(row.cells[1], center=True)
            
        if not tien_trinh:
            row = table.add_row()
            row.cells[0].text = "Lần 1: Nội dung cập nhật: "
            row.cells[1].text = ""

        # Ký tên
        doc.add_paragraph("\n")
        sig_table = doc.add_table(rows=2, cols=2)
        # Signature table should normally Not have borders, but docx default gives no borders unless style applied.
        sig_table.style = 'Normal Table'
        
        # force 100% width on signature
        from docx.oxml.shared import OxmlElement
        from docx.oxml.ns import qn
        tbl = sig_table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:type'), 'pct')
        tblW.set(qn('w:w'), '5000')
        tblPr.append(tblW)

        r0 = sig_table.rows[0]
        r1 = sig_table.rows[1]
        
        r0.cells[0].text = "TRƯỞNG KHOA"
        r1.cells[0].text = f"\n\n\n{data.get('ten_truong_khoa', '')}"
        
        r0.cells[1].text = f"Hà Nội, {data.get('ngay_ky', 'ngày ... tháng ... năm 20...')}\nNGƯỜI BIÊN SOẠN"
        r1.cells[1].text = f"\n\n\n{data.get('ten_nguoi_bien_soan', '')}"
        
        for c in r0.cells:
            apply_cell_style(c, bold=True, center=True)
        for c in r1.cells:
            apply_cell_style(c, bold=True, center=True)
