# services/word_builder/header_builder.py
from docx.shared import Pt, Inches
from templates.styles import apply_cell_style, set_table_style, COLOR_HEADER_BG, FONT_SIZE_NORMAL

class HeaderBuilder:
    @staticmethod
    def build(doc, data: dict):
        """Builds Section 1 and 2 of the syllabus."""
        HeaderBuilder._build_section_1(doc, data)
        HeaderBuilder._build_section_2(doc, data)

    @staticmethod
    def _build_section_1(doc, data: dict):
        # Heading 1
        h1 = doc.add_paragraph()
        run = h1.add_run("1. THÔNG TIN CHUNG VỀ HỌC PHẦN")
        run.font.bold = True
        run.font.size = FONT_SIZE_NORMAL

        HeaderBuilder._build_giang_vien_table(doc, data)
        doc.add_paragraph() # spacing
        HeaderBuilder._build_thong_tin_hp_table(doc, data)
        doc.add_paragraph() # spacing

    @staticmethod
    def _build_giang_vien_table(doc, data: dict):
        gv_chinh = data.get('giang_vien_chinh', [])
        gv_tham_gia = data.get('giang_vien_tham_gia', [])

        total_rows = 1 + len(gv_chinh) + (1 + len(gv_tham_gia) if gv_tham_gia else 0)
        table = doc.add_table(rows=total_rows, cols=6)
        set_table_style(table)

        # Build Title: Giảng viên phụ trách chính
        row_idx = 0
        r_chinh = table.rows[row_idx]
        apply_cell_style(r_chinh.cells[0], bold=True, bg_color=COLOR_HEADER_BG)
        # merge to end
        r_chinh.cells[0].merge(r_chinh.cells[5])
        r_chinh.cells[0].text = "Giảng viên phụ trách chính"
        
        row_idx += 1
        for i, gv in enumerate(gv_chinh):
            row = table.rows[row_idx]
            row.cells[0].text = str(gv.get('tt', i+1))
            row.cells[1].text = gv.get('hoc_ham_vi_ten', '')
            row.cells[2].text = gv.get('ma_can_bo', '')
            row.cells[3].text = gv.get('chuc_vu', '')
            row.cells[4].text = gv.get('sdt', '')
            row.cells[5].text = gv.get('email', '')
            apply_cell_style(row.cells[0], center=True)
            row_idx += 1

        if gv_tham_gia:
            r_tg = table.rows[row_idx]
            apply_cell_style(r_tg.cells[0], bold=True, bg_color=COLOR_HEADER_BG)
            r_tg.cells[0].merge(r_tg.cells[5])
            r_tg.cells[0].text = "Giảng viên tham gia giảng dạy"
            
            row_idx += 1
            for i, gv in enumerate(gv_tham_gia):
                row = table.rows[row_idx]
                row.cells[0].text = str(gv.get('tt', i+1))
                row.cells[1].text = gv.get('hoc_ham_vi_ten', '')
                row.cells[2].text = gv.get('ma_can_bo', '')
                row.cells[3].text = gv.get('chuc_vu', '')
                row.cells[4].text = gv.get('sdt', '')
                row.cells[5].text = gv.get('email', '')
                apply_cell_style(row.cells[0], center=True)
                row_idx += 1

    @staticmethod
    def _build_thong_tin_hp_table(doc, data: dict):
        table = doc.add_table(rows=8, cols=4)
        set_table_style(table, repeat_header=False)
        
        # Row 0: Tên học phần
        r0 = table.rows[0]
        r0.cells[0].text = "Tên học phần (tiếng Việt)"
        r0.cells[1].text = data.get('ten_tv', '')
        r0.cells[1].merge(r0.cells[3])
        
        # Row 1: Tên tiếng Anh
        r1 = table.rows[1]
        r1.cells[0].text = "Tên học phần (tiếng Anh)"
        r1.cells[1].text = data.get('ten_ta', '')
        r1.cells[1].merge(r1.cells[3])
        
        # Row 2: Mã học phần | Số tín chỉ
        r2 = table.rows[2]
        r2.cells[0].text = "Mã học phần"
        r2.cells[1].text = data.get('ma_hp', '')
        r2.cells[2].text = "Số tín chỉ"
        r2.cells[3].text = str(data.get('so_tc', ''))
        
        # Row 3: Loại học phần
        r3 = table.rows[3]
        r3.cells[0].text = "Loại học phần"
        r3.cells[1].text = data.get('loai_hp', '')
        r3.cells[1].merge(r3.cells[3])

        # Row 4: Tính chất
        r4 = table.rows[4]
        r4.cells[0].text = "Tính chất"
        r4.cells[1].text = data.get('tinh_chat', '')
        r4.cells[1].merge(r4.cells[3])

        # Row 5: Phân bổ thời gian
        r5 = table.rows[5]
        r5.cells[0].text = "Phân bổ thời gian"
        pb = data.get('phan_bo_gio', {})
        
        p = r5.cells[1].paragraphs[0]
        p.add_run(f"- Lý thuyết, Bài tập, Kiểm tra: {pb.get('ly_thuyet', 0)} (giờ)\n")
        p.add_run(f"- Thực hành, Thí nghiệm: {pb.get('thuc_hanh_tn', 0)} (giờ)\n")
        p.add_run(f"- Thảo luận có nội dung: {pb.get('thao_luan', 0)} (giờ)\n")
        p.add_run(f"- Tiểu luận, Đồ án: {pb.get('tieu_luan_do_an', 0)} (giờ)\n")
        p.add_run(f"- Thực tập tại DN: {pb.get('thuc_tap', 0)} (giờ)\n")
        p.add_run(f"- Tự học, tự NC: {pb.get('tu_hoc', 0)} (giờ)\n")
        tong = pb.get('ly_thuyet', 0) + pb.get('thuc_hanh_tn', 0) + pb.get('thao_luan', 0) + pb.get('tieu_luan_do_an', 0) + pb.get('thuc_tap', 0) + pb.get('tu_hoc', 0)
        p.add_run(f"- Tổng giờ học tập định mức: {tong} (giờ)").font.bold = True
        r5.cells[1].merge(r5.cells[3])

        # Row 6: Đơn vị quản lý
        r6 = table.rows[6]
        r6.cells[0].text = "Đơn vị quản lý học phần"
        r6.cells[1].text = data.get('don_vi', '')
        r6.cells[1].merge(r6.cells[3])

        # Row 7: Các HP liên quan
        r7 = table.rows[7]
        r7.cells[0].text = "Các học phần liên quan"
        p_rel = r7.cells[1].paragraphs[0]
        
        tq = data.get('hoc_phan_tien_quyet', [])
        p_rel.add_run("- Tiên quyết: ").font.bold = True
        p_rel.add_run(", ".join([f"{x['ten']} ({x['ma']})" for x in tq]) if tq else "Không\n")
        
        tt = data.get('hoc_phan_thay_the', {})
        if not p_rel.text.endswith('\n'): p_rel.add_run('\n')
        p_rel.add_run("- Thay thế: ").font.bold = True
        p_rel.add_run(f"{tt.get('ten')} ({tt.get('ma')})\n" if tt else "Không\n")
        
        sh = data.get('hoc_phan_song_hanh', {})
        p_rel.add_run("- Song hành: ").font.bold = True
        p_rel.add_run(f"{sh.get('ten')} ({sh.get('ma')})" if sh else "Không")
        
        r7.cells[1].merge(r7.cells[3])

    @staticmethod
    def _build_section_2(doc, data: dict):
        h2 = doc.add_paragraph()
        run = h2.add_run("2. MÔ TẢ TÓM TẮT NỘI DUNG HỌC PHẦN")
        run.font.bold = True
        run.font.size = FONT_SIZE_NORMAL

        # Đại học: Describe in italic
        trinh_do = data.get('trinh_do', 'Đại học')
        loai_hp = data.get('tinh_chat', '') # or loai_hp
        
        p = doc.add_paragraph()
        run_desc = p.add_run(data.get('mo_ta', '(Bổ sung mô tả tóm tắt nội dung)'))
        run_desc.font.italic = True
        
        if trinh_do == 'Tiến sĩ':
            if data.get('mo_ta_tieu_luan'):
                doc.add_paragraph().add_run(data['mo_ta_tieu_luan']).font.italic = True
            if data.get('mo_ta_chuyen_de'):
                doc.add_paragraph().add_run(data['mo_ta_chuyen_de']).font.italic = True
        
        doc.add_paragraph() # spacing
