# services/word_builder/checklist_builder.py
from docx.enum.text import WD_ALIGN_PARAGRAPH

class ChecklistBuilder:
    @staticmethod
    def build(doc, data: dict):
        """Thêm page break và render trang checklist cuối document."""
        doc.add_page_break()
        
        # Tiêu đề
        p_title = doc.add_paragraph()
        run_title = p_title.add_run("PHIẾU TỰ KIỂM TRA ĐỀ CƯƠNG CHI TIẾT HỌC PHẦN (CHECKLIST)")
        run_title.font.bold = True
        run_title.font.size = 14 * 12700 # 14pt (Pt is tricky direct assignment, use doc.styles or 14pt = 14*12700 EMUs inside but simpler is format)
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p_sub = doc.add_paragraph()
        run_sub = p_sub.add_run("(Dành cho Giảng viên biên soạn - Cập nhật cho năm học 2026)")
        run_sub.font.italic = True
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Tên học phần: {data.get('ten_tv', '')}       Mã HP: {data.get('ma_hp', '')}")
        doc.add_paragraph()

        # Build checklists
        ChecklistBuilder._add_section(doc, "I. KIỂM TRA HÌNH THỨC & QUY CÁCH (WIZARD TOUCH)", [
            "Font Times New Roman 12pt, giãn dòng 1.1–1.2",
            "Đánh số mục tự động, nhất quán",
            "Bảng: Repeat Header Rows + AutoFit Window",
            "Thông tin GV: đầy đủ học hàm, học vị, email, SĐT"
        ])
        
        ChecklistBuilder._add_section(doc, "II. CHUẨN ĐẦU RA (CLO)", [
            "Mỗi CLO dùng động từ đo lường được (thang Bloom)",
            "Không dùng 'Biết', 'Hiểu', 'Nắm vững'",
            "CLO đối chiếu trực tiếp với PLO tương ứng",
            "Cấp độ CLO phù hợp vị trí học phần trong CTĐT"
        ])

        ChecklistBuilder._add_section(doc, "III. NỘI DUNG & KẾ HOẠCH TỰ HỌC", [
            "Tổng giờ = số TC × 50 (1TC = 50 giờ chuẩn)",
            "Mỗi tuần/chương có nhiệm vụ tự học cụ thể",
            "Ít nhất 1 giáo trình chính + TLTK cập nhật 5–10 năm",
            "Có link học liệu số (nếu có)"
        ])

        ChecklistBuilder._add_section(doc, "IV. ĐÁNH GIÁ & RUBRICS", [
            "Bài KT bao phủ toàn bộ CLO đã công bố",
            "Hình thức thi phù hợp tính chất học phần",
            "Đã xây dựng Rubric chi tiết cho đánh giá phức tạp"
        ])

        ChecklistBuilder._add_section(doc, "V. CHÍNH SÁCH HỌC PHẦN (XU HƯỚNG 2026)", [
            "Đã nêu rõ quy định liêm chính học thuật & xử lý vi phạm",
            "Đã quy định mức độ cho phép dùng Generative AI"
        ])

        ChecklistBuilder._add_section(doc, "VI. XÁC NHẬN GV", [
            "Ký tên + ghi rõ họ tên + ngày tháng năm"
        ])

    @staticmethod
    def _add_section(doc, title: str, items: list):
        doc.add_paragraph(title).runs[0].font.bold = True
        for item in items:
            # Dùng ký tự unicode check box rỗng
            doc.add_paragraph(f"☐  {item}")
        doc.add_paragraph()
