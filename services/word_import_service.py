# services/word_import_service.py
import docx
import re

def import_dccthp(file_path: str) -> dict:
    """
    Parse file .docx ĐCCTHP có sẵn → trả về dict theo schema.
    Dùng python-docx đọc toàn bài và dùng regex matching để tách logic.
    """
    data = {
        "trinh_do": "Đại học",
        "ten_tv": "",
        "ten_ta": "",
        "don_vi": "",
        "ma_hp": "",
        "so_tc": 0,
        "loai_hp": "Bắt buộc",
        "tinh_chat": "Lý thuyết",
        "loai_hinh": "Trực tiếp",
        "phan_bo_gio": {
            "ly_thuyet": 0,
            "thuc_hanh_tn": 0,
            "thao_luan": 0,
            "tieu_luan_do_an": 0,
            "thuc_tap": 0,
            "tu_hoc": 0
        },
        "hoc_phan_tien_quyet": [],
        "hoc_phan_thay_the": {},
        "hoc_phan_song_hanh": {},
        "giang_vien_chinh": [],
        "giang_vien_tham_gia": [],
        "mo_ta": "",
        "mo_ta_tieu_luan": "",
        "mo_ta_chuyen_de": "",
        "muc_tieu": [],
        "clo": [],
        "giao_trinh": [],
        "tai_lieu_tk": [],
        "tai_lieu_khac": "",
        "phong_hoc": "",
        "thiet_bi": "",
        "thiet_bi_th": "",
        "bai_bao_quoc_te": [],
        "khong_gian_nc": "",
        "hoat_dong_ngoai_khoa": "",
        "noi_dung_chi_tiet": [],
        "thanh_phan_dg": [],
        "rubrics": [],
        "liem_chinh_nguong_plagiarism": 30,
        "ai_policy_level": 1,
        "ai_policy_mo_ta": "",
        "tien_trinh": [],
        "ten_truong_khoa": "",
        "ten_nguoi_bien_soan": "",
        "ngay_ky": ""
    }

    try:
        doc = docx.Document(file_path)
    except Exception:
        return data

    # Lấy text từng đoạn (Paragraphs parsing)
    full_text = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    # 1. Tên TV / Tên TA / Header fallback
    for text in full_text[:10]:
        # Nếu chuỗi thuần in hoa, độ dài > 10, không chứa các từ header chuẩn
        if text.isupper() and len(text) > 10 and "CỘNG HÒA" not in text and "TRƯỜNG ĐẠI HỌC" not in text and "ĐỀ CƯƠNG" not in text:
            if not data['ten_tv']:
                data['ten_tv'] = text

    # Parse Tables (Tùy thuộc index của Table)
    if len(doc.tables) > 1:
        # Table 0: header Đơn vị, Cộng hòa xã hội. 
        # Table 1: Giảng viên hoặc Thông tin học phần
        # Dựa trên Header Keyword để xác định
        for table in doc.tables:
            try:
                first_cell = table.cell(0, 0).text.lower()
                if "giảng viên phụ trách" in first_cell or "giảng viên tham gia" in first_cell:
                    _parse_gv_table(table, data)
                elif "tên học phần" in first_cell or "mã học phần" in first_cell:
                    _parse_thong_tin_hp(table, data)
                elif "mục tiêu" in first_cell and "mô tả" in first_cell:
                    _parse_muc_tieu(table, data)
                elif "cđr học phần" in first_cell or "clo" in first_cell:
                    _parse_clo(table, data)
                # ... Parse các bảng khác tương tự với logic index mapper ...
            except Exception:
                pass # Fallback nếu bảng bị lỗi

    return data

def _parse_gv_table(table, data):
    for row in table.rows[1:]:
        if len(row.cells) >= 4:
            gv = {
                "tt": row.cells[0].text.strip(),
                "hoc_ham_vi_ten": row.cells[1].text.strip(),
                "sdt": row.cells[2].text.strip(),
                "email": row.cells[3].text.strip()
            }
            if gv['hoc_ham_vi_ten']:
                # Mặc định thêm vào danh sách GV chính cho mẫu
                data['giang_vien_chinh'].append(gv)

def _parse_thong_tin_hp(table, data):
    # Rà soát từng ô, map key
    for row in table.rows:
        text0 = row.cells[0].text.lower()
        if "số tín chỉ" in text0 or (len(row.cells) > 2 and "số tín chỉ" in row.cells[2].text.lower()):
            try:
                val = row.cells[3].text.strip()
                data['so_tc'] = int(val)
            except Exception:
                pass
        if "tính chất" in text0:
            data['tinh_chat'] = row.cells[1].text.strip()

def _parse_muc_tieu(table, data):
     for row in table.rows[1:]:
        if len(row.cells) >= 3:
            mt = {
                "ma": row.cells[0].text.strip(),
                "mo_ta": row.cells[1].text.strip(),
                "plo": row.cells[2].text.strip()
            }
            if mt['ma']: data['muc_tieu'].append(mt)

def _parse_clo(table, data):
     for row in table.rows[1:]:
        if len(row.cells) >= 4:
            clo = {
                "ma": row.cells[0].text.strip(),
                "mo_ta": row.cells[1].text.strip(),
                "plo": row.cells[2].text.strip(),
                "muc_do": row.cells[3].text.strip()
            }
            if clo['ma']: data['clo'].append(clo)
