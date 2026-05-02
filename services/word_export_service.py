# services/word_export_service.py
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import os

from services.word_validator import validate_data, DCCTValidationError
from templates.styles import enforce_font_on_doc, FONT_SIZE_TITLE
from services.word_builder import (
    HeaderBuilder, CloBuilder, ContentBuilder, AssessmentBuilder,
    PolicyBuilder, ChecklistBuilder, SignatureBuilder
)

class DCCTExportError(Exception):
    pass

def export_dccthp(data: dict, output_path: str, template_type: str = "auto") -> str:
    """
    Export ĐCCTHP ra file .docx theo mẫu 2026.
    """
    # 1. Validate
    errors = validate_data(data)
    if errors:
        raise DCCTValidationError(errors)
    
    # 2. Setup Document
    # We create a new blank document. 
    # (If using word_template_*.docx, we could docx.Document("templates/word_template_..."))
    doc = docx.Document()
    
    # Configure global margin: trái 3cm, phải 2cm, trên 2cm, dưới 2cm
    # python-docx uses Cm for margins
    from docx.shared import Cm
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)
        
    try:
        # 3. Create document header (using 2-column table)
        tb_header = doc.add_table(rows=1, cols=2)
        tb_header.autofit = True
        c_left = tb_header.cell(0, 0)
        c_right = tb_header.cell(0, 1)
        
        p_l1 = c_left.paragraphs[0]
        p_l1.add_run("TRƯỜNG ĐẠI HỌC ĐIỆN LỰC\n").font.bold = False
        p_l1.add_run(f"KHOA {data.get('don_vi', '..............................').upper()}").font.bold = True
        p_l1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p_r1 = c_right.paragraphs[0]
        p_r1.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n").font.bold = True
        p_r1.add_run("Độc lập - Tự do - Hạnh phúc").font.bold = True
        p_r1.add_run("\n--------------------------------").font.bold = False
        p_r1.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph() # spacing

        # Title ĐCCTHP
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = p_title.add_run("ĐỀ CƯƠNG CHI TIẾT HỌC PHẦN")
        run_title.font.bold = True
        run_title.font.size = FONT_SIZE_TITLE
        
        # Tên Tiếng Việt
        p_tv = doc.add_paragraph()
        p_tv.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_tv = p_tv.add_run(data.get('ten_tv', '').upper())
        run_tv.font.bold = True
        run_tv.font.size = FONT_SIZE_TITLE
        
        # Tên Tiếng Anh
        p_ta = doc.add_paragraph()
        p_ta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if data.get('ten_ta'):
            run_ta = p_ta.add_run(f"({data.get('ten_ta', '')})")
            run_ta.font.italic = True
        doc.add_paragraph()

        # 4. Call builders
        HeaderBuilder.build(doc, data)
        CloBuilder.build(doc, data)
        ContentBuilder.build(doc, data)
        AssessmentBuilder.build(doc, data)
        PolicyBuilder.build(doc, data)
        SignatureBuilder.build(doc, data)
        ChecklistBuilder.build(doc, data)

        # 5. Apply global styles
        enforce_font_on_doc(doc)
        
        # Paragraph spacing 1.15 lines 
        for paragraph in doc.paragraphs:
            paragraph.paragraph_format.line_spacing = 1.15
            
        # 6. Save
        doc.save(output_path)
        return output_path
        
    except Exception as e:
        raise DCCTExportError(f"Export Error: {str(e)}")
