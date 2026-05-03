import unittest
import os
import tempfile
import shutil
import docx

from services.word_export_service import export_dccthp
from services.word_import_service import import_dccthp
from services.word_validator import validate_data, DCCTValidationError

# Giả lập dữ liệu chuẩn
def valid_base_data():
    return {
        "trinh_do": "Đại học",
        "ten_tv": "HỌC PHẦN MẪU",
        "ten_ta": "SAMPLE COURSE",
        "don_vi": "KHOA ĐÀO TẠO ĐẠI CƯƠNG",
        "ma_hp": "DC101",
        "so_tc": 3,
        "loai_hp": "Bắt buộc",
        "tinh_chat": "Lý thuyết",
        "loai_hinh": "Trực tiếp",
        "phan_bo_gio": {"ly_thuyet": 30, "tu_hoc": 120},
        "giang_vien_chinh": [{"tt": 1, "hoc_ham_vi_ten": "TS. A"}],
        "clo": [{"ma": "CLO1", "mo_ta": "Phân tích được abc...", "muc_do": "R"}],
        "thanh_phan_dg": [
            {"thanh_phan": "Giữa kỳ", "trong_so": 30},
            {"thanh_phan": "Cuối kỳ", "trong_so": 70}
        ],
        "rubrics": [{"ma": "R1", "tieu_chi": [{"ten": "T1", "trong_so": 100}]}],
        "ai_policy_level": 1,
    }

class TestWordServices(unittest.TestCase):
    def setUp(self):
        self.test_dir = tempfile.mkdtemp()

    def tearDown(self):
        shutil.rmtree(self.test_dir)

    def test_t5_validation_clo_hieu(self):
        data = valid_base_data()
        data['clo'][0]['mo_ta'] = "Hiểu các quy luật cơ bản"
        with self.assertRaises(DCCTValidationError) as context:
            export_dccthp(data, os.path.join(self.test_dir, "test.docx"))
        self.assertIn("Biết/Hiểu/Nắm vững", str(context.exception))

    def test_t6_validation_tong_gio(self):
        data = valid_base_data()
        data['phan_bo_gio']['tu_hoc'] = 100  # tổng = 130 != 150
        with self.assertRaises(DCCTValidationError) as context:
            export_dccthp(data, os.path.join(self.test_dir, "test.docx"))
        self.assertIn("Tổng giờ", str(context.exception))

    def test_t1_export_lt_dh(self):
        data = valid_base_data()
        output_path = os.path.join(self.test_dir, "t1_lt.docx")
        export_dccthp(data, output_path)
        self.assertTrue(os.path.exists(output_path))
        self.assertGreater(os.path.getsize(output_path), 10000)

    def test_t2_export_hon_hop_dh(self):
        data = valid_base_data()
        data['tinh_chat'] = "Hỗn hợp"
        data['phan_bo_gio'] = {"ly_thuyet": 10, "thuc_hanh_tn": 10, "tu_hoc": 130}
        data['noi_dung_chi_tiet'] = [
            {"tuan": "1", "noi_dung": "LT Bài 1", "gio_lt": "(5/0/0/0)"},
            {"tuan": "2", "noi_dung": "TH Bài 1", "gio_lt": "(0/0/0/TH)"}
        ]
        output_path = os.path.join(self.test_dir, "t2_hh.docx")
        export_dccthp(data, output_path)
        self.assertTrue(os.path.exists(output_path))
        
        doc = docx.Document(output_path)
        texts = [p.text for p in doc.paragraphs]
        self.assertTrue(any("6.1. Nội dung lý thuyết" in t for t in texts))
        self.assertTrue(any("6.2. Nội dung thực hành/thí nghiệm" in t for t in texts))

    def test_t3_export_tieu_luan_ts(self):
        data = valid_base_data()
        data['trinh_do'] = "Tiến sĩ"
        data['tinh_chat'] = "Tiểu luận tổng quan"
        data['mo_ta_tieu_luan'] = "Review lý thuyết chuyên ngành"
        
        output_path = os.path.join(self.test_dir, "t3_ts.docx")
        export_dccthp(data, output_path)
        
        doc = docx.Document(output_path)
        texts = [p.text for p in doc.paragraphs]
        self.assertTrue(any("Tiểu luận tổng quan" in t for t in texts))

    def test_t4_export_chuyen_de_ts(self):
        data = valid_base_data()
        data['trinh_do'] = "Tiến sĩ"
        data['tinh_chat'] = "Chuyên đề Tiến sĩ"
        
        output_path = os.path.join(self.test_dir, "t4_ts.docx")
        export_dccthp(data, output_path)
        self.assertTrue(os.path.exists(output_path))

    def test_t7_import_sample_word(self):
        data = valid_base_data()
        output_path = os.path.join(self.test_dir, "sample.docx")
        export_dccthp(data, output_path)
        
        imported = import_dccthp(output_path)
        self.assertIn("trinh_do", imported)
        self.assertIn("mo_ta", imported)

    def test_t8_export_import_diff(self):
        data = valid_base_data()
        output_path = os.path.join(self.test_dir, "t8_diff.docx")
        export_dccthp(data, output_path)
        
        imported = import_dccthp(output_path)
        self.assertEqual(imported['ten_tv'].upper(), data['ten_tv'])

    def test_t9_rubric_table(self):
        data = valid_base_data()
        output_path = os.path.join(self.test_dir, "t9_rubric.docx")
        export_dccthp(data, output_path)
        
        doc = docx.Document(output_path)
        headers = [c.text for t in doc.tables for r in t.rows for c in r.cells]
        self.assertIn("Xuất sắc (9–10)", headers)
        
    def test_t10_checklist_page(self):
        data = valid_base_data()
        output_path = os.path.join(self.test_dir, "t10_checklist.docx")
        export_dccthp(data, output_path)
        
        doc = docx.Document(output_path)
        texts = [p.text for p in doc.paragraphs]
        self.assertTrue(any("PHIẾU TỰ KIỂM TRA ĐỀ CƯƠNG" in t for t in texts))
        self.assertTrue(any("VI. XÁC NHẬN GV" in t for t in texts))

if __name__ == '__main__':
    unittest.main()
