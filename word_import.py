# word_import.py - Pure Parser
import re
from docx import Document
import unicodedata

class DCCTHPImporter:
    """Parser class for Syllabus Word documents. 
    NO DATABASE LOGIC ALLOWED HERE."""
    def __init__(self):
        self.data = {
            'hp': {},
            'giang_vien': [],
            'muc_tieu': [],
            'clos': [],
            'hoc_lieu': [],
            'noi_dung': [],
            'danh_gia': [],
            'rubrics': []
        }
        self.logs = []


    def import_from_docx(self, file_path):
        try:
            doc = Document(file_path)

            # 1. Parse Header/Paragraphs for HP Info
            try:
                self._parse_paragraphs(doc.paragraphs)
                if self.data['hp'].get('ten_viet'):
                    self.logs.append(('success', 'Thông tin chung', 'Bóc tách thành công thông tin tiêu đề.'))
            except Exception as e:
                self.logs.append(('error', 'Thông tin chung', f'Lỗi paragraph: {str(e)}'))

            # 2. Iterate Tables — dùng logic nhận diện cải tiến
            for i, table in enumerate(doc.tables):
                try:
                    self._process_table(table, i)
                except Exception as e:
                    self.logs.append(('error', f'Bảng số {i + 1}', f'Lỗi bóc tách: {str(e)}'))

            return {
                'data': self.data,
                'logs': self.logs
            }
        except Exception as e:
            self.logs.append(('error', 'Hệ thống', f'Không thể mở file Word: {str(e)}'))
            return {'data': self.data, 'logs': self.logs}

    def _clean(self, text):
        if not text: return ""
        text = str(text).strip()
        # NFC normalization for Vietnamese
        return unicodedata.normalize('NFC', text)

    def _cell_text(self, cell):
        return self._clean(" ".join(p.text for p in cell.paragraphs))

    def _parse_paragraphs(self, paras):
        for p in paras:
            text = self._clean(p.text)
            if not text: continue

            # HP name follows ĐỀ CƯƠNG CHI TIẾT or similar
            if "HỌC PHẦN:" in text.upper():
                self.data['hp']['ten_viet'] = text.split(":")[-1].strip()

            # Parse Vietnamese name from uppercase bold text (title)
            # Pattern: all-caps line after "ĐỀ CƯƠNG CHI TIẾT HỌC PHẦN"
            if p.runs and all(r.bold for r in p.runs if r.text.strip()):
                txt = text.strip()
                if txt.isupper() and len(txt) > 5 and not self.data['hp'].get('ten_viet'):
                    self.data['hp']['ten_viet'] = txt.title()

            # Trình độ
            if "Trình độ đào tạo:" in text:
                self.data['hp']['trinh_do'] = text.split(":")[-1].strip()

    def _process_table(self, table, table_idx=0):
        """FIXED M-10: Cải thiện logic nhận diện bảng, tránh nhầm lẫn."""
        if not table.rows:
            return

        # Lấy text tất cả cells trong 2 hàng đầu để nhận diện tốt hơn
        header_cells = []
        for row_i in range(min(2, len(table.rows))):
            for cell in table.rows[row_i].cells:
                header_cells.append(self._cell_text(cell).lower())
        header_text = " | ".join(header_cells)

        # Thứ tự nhận diện quan trọng: từ đặc thù nhất đến chung nhất
        # 1. Rubric — đặc thù nhất
        if ("mức xuất sắc" in header_text or "muc xuat sac" in header_text or "xuất sắc" in header_text) \
                and ("tiêu chí" in header_text or "tieu chi" in header_text):
            self._parse_rubric(table)
            self.logs.append(('success', 'Rubric', 'Bóc tách xong bảng Rubric.'))

        # 2. CLO — có "chuẩn đầu ra" và "mô tả"
        elif ("chuẩn đầu ra" in header_text or "cdr học phần" in header_text or "clo" in header_text) \
                and ("mô tả" in header_text or "mo ta" in header_text):
            self._parse_clo(table)
            self.logs.append(('success', 'CLO', f'Đã nhập {len(self.data["clos"])} chuẩn đầu ra.'))

        # 3. Bảng đánh giá — có "thành phần đánh giá" hoặc "trọng số" + "hình thức"
        elif ("thành phần đánh giá" in header_text or "thanh phan danh gia" in header_text) \
                or ("trọng số" in header_text and ("hình thức" in header_text or "hinh thuc" in header_text)):
            self._parse_danh_gia(table)
            self.logs.append(('success', 'Đánh giá', f'Đã nhập {len(self.data["danh_gia"])} bài đánh giá.'))

        # 4. Bảng nội dung chi tiết — có "nội dung" và "giờ" (lt/bt/th)
        elif ("nội dung" in header_text or "noi dung" in header_text) \
                and ("giờ" in header_text or "gio" in header_text or "lt" in header_text or "bt" in header_text):
            self._parse_noi_dung(table)
            self.logs.append(('success', 'Nội dung', f'Đã nhập {len(self.data["noi_dung"])} mục nội dung.'))

        # 5. Mục tiêu — có "mục tiêu" và "mô tả"
        elif ("mục tiêu" in header_text or "muc tieu" in header_text) \
                and "mô tả" in header_text:
            self._parse_muc_tieu(table)
            self.logs.append(('success', 'Mục tiêu', f'Đã nhập {len(self.data["muc_tieu"])} mục tiêu.'))

        # 6. Bảng GV — phải kiểm tra sau CLO để tránh nhầm
        elif ("học hàm" in header_text or "hoc ham" in header_text
              or ("số điện thoại" in header_text and "email" in header_text)):
            self._parse_giang_vien(table)
            self.logs.append(('success', 'Giảng viên', f'Đã nhập {len(self.data["giang_vien"])} giảng viên.'))

        # 7. Bảng thông tin HP (mã HP, tín chỉ, loại) — rất chung, kiểm tra sau cùng
        elif "mã học phần" in header_text or ("số tín chỉ" in header_text and "loại học phần" in header_text):
            self._parse_hp_table(table)
            self.logs.append(('success', 'Thông tin mã HP', 'Đã cập nhật mã và số tín chỉ.'))

        # 8. Bảng tài liệu — kiểm tra sau DG để tránh nhầm với rubric
        elif "tài liệu" in header_text and ("tác giả" in header_text or "năm" in header_text or "nhà xuất bản" in header_text):
            self._parse_tai_lieu(table)
            self.logs.append(('success', 'Học liệu', f'Đã nhập {len(self.data["hoc_lieu"])} tài liệu.'))

    def _parse_hp_table(self, table):
        """Parse bảng thông tin chung HP."""
        for row in table.rows:
            cells = [self._cell_text(c) for c in row.cells]
            for i, c in enumerate(cells):
                c_low = c.lower()
                if "mã học phần" in c_low and i + 1 < len(cells):
                    self.data['hp']['ma'] = cells[i + 1]
                elif "số tín chỉ" in c_low and i + 1 < len(cells):
                    m = re.search(r'(\d+)', cells[i + 1])
                    if m: self.data['hp']['so_tin_chi'] = int(m.group(1))
                elif "loại học phần" in c_low and i + 1 < len(cells):
                    self.data['hp']['loai'] = cells[i + 1]
                elif "tính chất" in c_low and i + 1 < len(cells):
                    self.data['hp']['tinh_chat'] = cells[i + 1]
                elif "đơn vị quản lý" in c_low and i + 1 < len(cells):
                    self.data['hp']['don_vi_ql'] = cells[i + 1]
                elif "loại hình" in c_low and i + 1 < len(cells):
                    self.data['hp']['loai_hinh'] = cells[i + 1]
                elif "tổng giờ" in c_low and i + 1 < len(cells):
                    m = re.search(r'(\d+)', cells[i + 1])
                    if m: self.data['hp']['tong_gio'] = int(m.group(1))
                elif "tiên quyết" in c_low and i + 1 < len(cells):
                    self.data['hp']['hp_tien_quyet'] = cells[i + 1]
                elif "song hành" in c_low and i + 1 < len(cells):
                    self.data['hp']['hp_song_hanh'] = cells[i + 1]
                elif "thay thế" in c_low and i + 1 < len(cells):
                    self.data['hp']['hp_thay_the'] = cells[i + 1]

    def _parse_giang_vien(self, table):
        """Parse bảng giảng viên."""
        for row in table.rows[1:]:  # Bỏ header
            cells = [self._cell_text(c) for c in row.cells]
            if len(cells) >= 2 and cells[1]:  # Cột thứ 2 là Họ tên
                self.data['giang_vien'].append({
                    'ho_ten': cells[1] if len(cells) > 1 else '',
                    'sdt': cells[2] if len(cells) > 2 else '',
                    'email': cells[3] if len(cells) > 3 else '',
                    'vai_tro': 'tham_gia'
                })

    def _parse_muc_tieu(self, table):
        """Parse bảng mục tiêu."""
        for i, row in enumerate(table.rows[1:]):
            cells = [self._cell_text(c) for c in row.cells]
            if len(cells) >= 2 and cells[1]:
                self.data['muc_tieu'].append({
                    'so_thu_tu': i + 1,
                    'mo_ta': cells[1],
                    'cdr_ma': cells[2] if len(cells) > 2 else ''
                })

    def _parse_clo(self, table):
        """FIXED C-01: Dùng key 'ma' thay vì 'ma_clo' cho bảng production."""
        for row in table.rows[1:]:
            cells = [self._cell_text(c) for c in row.cells]
            if len(cells) >= 2 and cells[0]:
                # FIXED: Dùng key 'ma' — khớp với cột trong bảng production 'clo'
                self.data['clos'].append({
                    'ma': cells[0],                          # FIXED: 'ma' không phải 'ma_clo'
                    'mo_ta': cells[1] if len(cells) > 1 else '',
                    'cdr_ma': cells[2] if len(cells) > 2 else '',
                    'level_irm': cells[3] if len(cells) > 3 else 'I'
                })

    def _parse_tai_lieu(self, table):
        """Parse bảng tài liệu."""
        for row in table.rows[1:]:
            cells = [self._cell_text(c) for c in row.cells]
            if len(cells) >= 1 and cells[0]:
                full_text = " ".join(c for c in cells if c).strip()
                if not full_text:
                    continue
                # Cố gắng nhận biết loại tài liệu từ nội dung hoặc context
                loai = 'tham_khao'
                if any(kw in full_text.lower() for kw in ['giáo trình', 'sách giáo khoa', 'tài liệu học tập']):
                    loai = 'chinh'
                self.data['hoc_lieu'].append({
                    'loai': loai,
                    'noi_dung': full_text,
                    'so_thu_tu': len(self.data['hoc_lieu']) + 1
                })

    def _parse_noi_dung(self, table):
        """Parse bảng nội dung chi tiết."""
        idx = 1
        for row in table.rows[1:]:
            cells = [self._cell_text(c) for c in row.cells]
            if not any(cells): continue

            # Nhận biết hàng tiêu đề chương (bold hoặc không có số ở cột 0)
            is_chapter = not cells[0].strip().isdigit() and not cells[0].strip() == ''
            ten = cells[1] if len(cells) > 1 else cells[0]
            if not ten: continue

            # Parse giờ — định dạng phổ biến: "2/0/0/0" hoặc riêng lẻ từng cột
            gio_lt = gio_bt = gio_tl = gio_th = 0
            gio_cell = cells[2] if len(cells) > 2 else ''
            if '/' in gio_cell:
                parts = gio_cell.split('/')
                try:
                    gio_lt = int(re.search(r'\d+', parts[0]).group()) if parts[0].strip() else 0
                    gio_bt = int(re.search(r'\d+', parts[1]).group()) if len(parts) > 1 and parts[1].strip() else 0
                    gio_tl = int(re.search(r'\d+', parts[2]).group()) if len(parts) > 2 and parts[2].strip() else 0
                    gio_th = int(re.search(r'\d+', parts[3]).group()) if len(parts) > 3 and parts[3].strip() else 0
                except (AttributeError, ValueError):
                    pass
            else:
                m = re.search(r'\d+', gio_cell)
                if m: gio_lt = int(m.group())

            self.data['noi_dung'].append({
                'phan': 'lt',
                'thu_tu': idx,
                'ten': ten,
                'loai': 'chuong' if is_chapter else 'thuong',
                'cap_do': 1 if is_chapter else 2,
                'gio_lt': gio_lt,
                'gio_bt': gio_bt,
                'gio_tl': gio_tl,
                'gio_th_tn': gio_th,
                'pp_day': cells[3] if len(cells) > 3 else '',
                'pp_hoc': cells[4] if len(cells) > 4 else '',
                'cdr_ma': cells[5] if len(cells) > 5 else '',
                'bai_danh_gia': cells[6] if len(cells) > 6 else ''
            })
            if not is_chapter:
                idx += 1

    def _parse_danh_gia(self, table):
        """Parse bảng kế hoạch kiểm tra - đánh giá."""
        for row in table.rows[1:]:
            cells = [self._cell_text(c) for c in row.cells]
            if not cells[0] and not cells[1]: continue

            nhom = cells[0] if cells[0] else 'thuong_xuyen'
            # Normalize nhóm
            if any(kw in nhom.lower() for kw in ['thường xuyên', 'thuong xuyen', 'quá trình']):
                nhom_key = 'thuong_xuyen'
            elif any(kw in nhom.lower() for kw in ['cuối kỳ', 'cuoi ky', 'kỳ thi']):
                nhom_key = 'cuoi_ky'
            else:
                nhom_key = nhom

            # Trọng số nhóm — thường ở cột 1
            ty_trong_nhom = 0.0
            if len(cells) > 1:
                m = re.search(r'(\d+(?:\.\d+)?)', cells[1].replace('%', ''))
                if m:
                    val = float(m.group(1))
                    ty_trong_nhom = val if val <= 1.1 else val  # giữ nguyên dạng % hay ratio

            self.data['danh_gia'].append({
                'nhom': nhom_key,
                'ty_trong_nhom': ty_trong_nhom,
                'noi_dung': cells[2] if len(cells) > 2 else '',
                'hinh_thuc': cells[3] if len(cells) > 3 else '',
                'tieu_chi_danh_gia': cells[4] if len(cells) > 4 else '',
                'clo_lien_quan': cells[5] if len(cells) > 5 else '',
                'diem_toi_da_cdr': cells[6] if len(cells) > 6 else '10',
                'trong_so_cdr': cells[7] if len(cells) > 7 else '',
                'thu_tu': len(self.data['danh_gia']) + 1
            })

    def _parse_rubric(self, table):
        """Parse bảng rubric: Tiêu chí | Trọng số | Xuất sắc | Tốt | Đạt | Chưa đạt."""
        rubric_item = {
            'ten': 'Rubric đánh giá',
            'ky_hieu': f"R{len(self.data['rubrics']) + 1}",
            'mo_ta': '',
            'thu_tu': len(self.data['rubrics']) + 1,
            'tieu_chi_list': []
        }

        for row in table.rows[1:]:
            cells = [self._cell_text(c) for c in row.cells]
            if len(cells) >= 2 and cells[0]:
                rubric_item['tieu_chi_list'].append({
                    'tieu_chi': cells[0],
                    'trong_so': cells[1] if len(cells) > 1 else '',
                    'muc_xuat_sac': cells[2] if len(cells) > 2 else '',
                    'muc_tot': cells[3] if len(cells) > 3 else '',
                    'muc_dat': cells[4] if len(cells) > 4 else '',
                    'muc_chua_dat': cells[5] if len(cells) > 5 else '',
                    'thu_tu': len(rubric_item['tieu_chi_list']) + 1
                })

        if rubric_item['tieu_chi_list']:
            self.data['rubrics'].append(rubric_item)


def parse_docx(file_path):
    """Wrapper function for service compatibility."""
    importer = DCCTHPImporter()
    return importer.import_from_docx(file_path)


# Removed legacy import_single function as per P2 refactor. 
# Logic moved to ImportExportService.save_imported_data()

