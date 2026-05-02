from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import os

def set_cell_background(cell, fill_color):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_modern_template(path):
    doc = Document()
    # Margins
    sec = doc.sections[0]
    sec.left_margin, sec.right_margin = Cm(3), Cm(2)
    sec.top_margin, sec.bottom_margin = Cm(2), Cm(2)

    # Header
    p = doc.add_paragraph('TRƯỜNG ĐẠI HỌC ĐIỆN LỰC')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph('{{ Department | upper }}').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n\n')
    
    # Title
    t = doc.add_paragraph('ĐỀ CƯƠNG CHI TIẾT HỌC PHẦN')
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run('\n{{ CourseName | upper }}')
    tr.bold = True
    tr.font.size = Pt(18)
    tr.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph('Mã học phần: {{ CourseCode }}').alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Sec 1: Info
    h1 = doc.add_paragraph('1. THÔNG TIN CHUNG')
    h1.runs[0].bold = True
    h1.runs[0].font.size = Pt(13)
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    rows = table.rows
    rows[0].cells[0].text = 'Số tín chỉ:'
    rows[0].cells[1].text = '{{ Credits }}'
    rows[1].cells[0].text = 'Loại học phần:'
    rows[1].cells[1].text = '{{ CourseType }}'
    rows[2].cells[0].text = 'Trình độ đào tạo:'
    rows[2].cells[1].text = '{{ Level }}'
    rows[3].cells[0].text = 'Đơn vị quản lý:'
    rows[3].cells[1].text = '{{ ManageUnit }}'

    # Sec 2: CLOs
    doc.add_paragraph('\n2. CHUẨN ĐẦU RA (CLO)')
    clo_table = doc.add_table(rows=2, cols=3)
    clo_table.style = 'Table Grid'
    hdr = clo_table.rows[0].cells
    hdr[0].text = 'Mã CLO'
    hdr[1].text = 'Mô tả chi tiết'
    hdr[2].text = 'CĐR CTĐT (PLO)'
    for c in hdr: set_cell_background(c, 'D9EAD3')
    
    # Loop
    r = clo_table.rows[1].cells
    r[0].text = '{% for clo in CLOs %}{{ clo.Code }}'
    r[1].text = '{{ clo.Desc }}'
    r[2].text = '{{ clo.PLO }}{% if not loop.last %}\n\n{% endif %}{% endfor %}'

    # Footer
    p_footer = doc.add_paragraph('\n\nNgày xuất: {{ Today }}')
    p_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.save(path)

def create_compact_template(path):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    
    doc.add_paragraph('EPU - SYLLABUS - {{ CourseCode }}').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    t = doc.add_paragraph('COURSE SPECIFICATION: {{ CourseName }}')
    t.runs[0].bold = True
    t.runs[0].font.size = Pt(14)
    
    doc.add_paragraph('Credits: {{ Credits }} | Dept: {{ Department }}')
    
    doc.add_paragraph('---')
    doc.add_paragraph('FACULTY TEAM:').runs[0].bold = True
    p = doc.add_paragraph('{% for lect in Lecturers %}')
    p.add_run('• {{ lect.Degree }} {{ lect.Name }}').bold = True
    p.add_run(' ({{ lect.Role }})\n  Email: {{ lect.Email }}\n{% endfor %}')
    
    doc.add_paragraph('---')
    doc.add_paragraph('DETAILED CONTENT:').runs[0].bold = True
    p2 = doc.add_paragraph('{% for nd in ContentLT %}')
    p2.add_run('Chương {{ loop.index }}: {{ nd.ten }}').bold = True
    p2.add_run(' - {{ nd.gio_lt }} tiết\n{% endfor %}')
    
    doc.save(path)

def create_quality_template(path):
    doc = Document()
    doc.add_paragraph('Hệ thống Đảm bảo Chất lượng - EPU').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\nPHÂN TÍCH CHUẨN ĐẦU RA & ĐÁNH GIÁ').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Học phần: {{ CourseName }} ({{ CourseCode }})').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n1. Ma trận CLO - PLO:').runs[0].bold = True
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'
    h = table.rows[0].cells
    h[0].text = 'CLO'
    h[1].text = 'PLO'
    h[2].text = 'Mức độ (IRM)'
    
    r = table.rows[1].cells
    r[0].text = '{% for clo in CLOs %}{{ clo.Code }}'
    r[1].text = '{{ clo.PLO }}'
    r[2].text = '{{ clo.Level }}{% if not loop.last %}\n{% endif %}{% endfor %}'

    doc.add_paragraph('\n2. Kế hoạch đánh giá:').runs[0].bold = True
    kt_table = doc.add_table(rows=2, cols=4)
    kt_table.style = 'Table Grid'
    h = kt_table.rows[0].cells
    h[0].text = 'Thành phần'
    h[1].text = 'Trọng số'
    h[2].text = 'Bài đánh giá'
    h[3].text = 'Hình thức'
    
    r = kt_table.rows[1].cells
    r[0].text = '{% for row in AssessmentRows %}{{ row.nhom }}'
    r[1].text = '{{ row.ty_trong }}%'
    r[2].text = '{{ row.noi_dung }}'
    r[3].text = '{{ row.hinh_thuc }}{% if not loop.last %}\n{% endif %}{% endfor %}'

    doc.save(path)

if __name__ == '__main__':
    target_dir = os.path.join(os.getcwd(), 'word_templates')
    os.makedirs(target_dir, exist_ok=True)
    
    create_modern_template(os.path.join(target_dir, 'Template_Hien_Dai.docx'))
    create_compact_template(os.path.join(target_dir, 'Template_Nhanh_Gon.docx'))
    create_quality_template(os.path.join(target_dir, 'Template_Kiem_Dinh_Chat_Luong.docx'))
    print("Done generating 3 test templates.")
