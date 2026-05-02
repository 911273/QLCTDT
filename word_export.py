# word_export.py
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

class DCCTHPExporter:
    def __init__(self, db):
        self.db = db
        self.doc = None

    def export(self, hp_id_or_ma, output_path):
        # Handle both ID and ma
        if isinstance(hp_id_or_ma, int) or (isinstance(hp_id_or_ma, str) and hp_id_or_ma.isdigit()):
             hp_row = self.db.conn.execute("SELECT * FROM hoc_phan WHERE id=?", (int(hp_id_or_ma),)).fetchone()
        else:
             hp_row = self.db.conn.execute("SELECT * FROM hoc_phan WHERE ma=? OR ma_hp=?", (hp_id_or_ma, hp_id_or_ma)).fetchone()
             
        if not hp_row: raise ValueError(f"Không tìm thấy học phần với định danh: {hp_id_or_ma}")
        hp = dict(hp_row)

        self.doc = Document()
        self._setup_page()
        
        # 1. Header & Title
        self._add_header_section(hp)
        
        # 2. Main Sections (Theo Mẫu 13 - EPU)
        self._add_sec1_thong_tin_chung(hp); self._append_extra_fields(hp, 'sec1_thong_tin_chung')
        self._add_sec2_mo_ta(hp); self._append_extra_fields(hp, 'sec2_mo_ta')
        self._add_sec3_muc_tieu(hp); self._append_extra_fields(hp, 'sec3_muc_tieu')
        self._add_sec4_clo(hp); self._append_extra_fields(hp, 'sec4_clo')
        self._add_sec5_hoc_lieu(hp); self._append_extra_fields(hp, 'sec5_hoc_lieu')
        self._add_sec6_noi_dung(hp); self._append_extra_fields(hp, 'sec6_noi_dung')
        self._add_sec7_pp_day_hoc(hp); self._append_extra_fields(hp, 'sec7_pp_day_hoc')
        self._add_sec8_kiem_tra_danh_gia(hp); self._append_extra_fields(hp, 'sec8_kiem_tra_danh_gia')
        self._add_sec9_quy_dinh(hp); self._append_extra_fields(hp, 'sec9_quy_dinh')
        self._add_sec10_co_so_phuc_vu(hp); self._append_extra_fields(hp, 'sec10_co_so_phuc_vu')
        self._add_sec11_doi_ngu_gv(hp); self._append_extra_fields(hp, 'sec11_doi_ngu_gv')
        self._add_sec12_phu_luc(hp); self._append_extra_fields(hp, 'sec12_phu_luc')
        self._add_sec13_cap_nhat(hp); self._append_extra_fields(hp, 'sec13_cap_nhat')
        self._add_sec14_chinh_sach(hp); self._append_extra_fields(hp, 'sec14_chinh_sach')
        self._add_sec15_checklist(hp); self._append_extra_fields(hp, 'sec15_checklist')
        
        # 3. Dynamic Sections from User (Bổ sung linh hoạt)
        dynamic_secs = self.db.list_sections()
        for ds in dynamic_secs:
            self._add_heading(ds['label'])
            self._append_extra_fields(hp, ds['section_key'])
            
        self._add_signatures(hp)
        
        self.doc.save(output_path)
        return True

    def _append_extra_fields(self, hp, section_key):
        """Tự động chèn các trường bổ sung (Inline CRUD) vào bản xuất Built-in."""
        try:
            extra_vals = self.db.load_extra_data(hp['id'], section_key)
            if not extra_vals: return
            
            # Lấy tên nhãn của các trường
            field_defs = {f['field_key']: f['nhan'] for f in self.db.list_extra_fields(section_key)}
            import json
            
            for fkey, val in extra_vals.items():
                if not val: continue
                label = field_defs.get(fkey, fkey)
                
                # Parse if JSON (Table data)
                data_obj = None
                if isinstance(val, str) and (val.startswith('[') or val.startswith('{')):
                    try: data_obj = json.loads(val)
                    except: pass
                
                if isinstance(data_obj, list): # Bảng dữ liệu
                    if data_obj:
                        p = self.doc.add_paragraph()
                        p.add_run(f"{label}:").font.bold = True
                        p.paragraph_format.space_before = Pt(6)
                        
                        cols = list(data_obj[0].keys())
                        tab = self.doc.add_table(rows=1, cols=len(cols))
                        tab.style = 'Table Grid'
                        for i, c in enumerate(cols):
                            cell = tab.rows[0].cells[i]
                            cell.text = c
                            cell.paragraphs[0].runs[0].font.bold = True
                            
                        for row_data in data_obj:
                            r = tab.add_row()
                            for i, c in enumerate(cols):
                                r.cells[i].text = str(row_data.get(c, ''))
                    continue

                # Text bình thường
                p = self.doc.add_paragraph()
                p.add_run(f"{label}: ").font.bold = True
                p.add_run(str(val))
        except Exception as e:
            print(f"[WordExport] Error appending extra fields for {section_key}: {e}")

    def _setup_page(self):
        # Font & Spacing
        style = self.doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        style.paragraph_format.line_spacing = 1.15
        
        # Margins: Left 3cm, Others 2cm
        for sec in self.doc.sections:
            sec.top_margin = Cm(2)
            sec.bottom_margin = Cm(2)
            sec.left_margin = Cm(3)
            sec.right_margin = Cm(2)

    def _add_header_section(self, hp):
        # FIXED N-02: Lấy tên khoa từ DB thay vì hardcode placeholder
        ten_khoa = hp.get('ten_khoa') or ''
        if not ten_khoa and hp.get('khoa_id'):
            k = self.db.conn.execute("SELECT ten FROM khoa WHERE id=?", (hp['khoa_id'],)).fetchone()
            if k: ten_khoa = k['ten']

        # Header Table (Logo | School Name)
        table = self.doc.add_table(rows=1, cols=2)
        table.width = Cm(17.5)

        # Left side: School info
        cell_left = table.rows[0].cells[0]
        p1 = cell_left.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run0 = p1.add_run("TRƯỜNG ĐẠI HỌC ĐIỆN LỰC")
        run0.font.bold = True
        run0.font.size = Pt(12)
        p1a = cell_left.add_paragraph(f"KHOA {ten_khoa.upper()}" if ten_khoa else "KHOA")
        p1a.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1a.runs[0].font.bold = True

        # Right side: Motto
        cell_right = table.rows[0].cells[1]
        p2 = cell_right.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\nĐộc lập - Tự do - Hạnh phúc")
        run2.font.bold = True
        run2.font.size = Pt(11)
        p2a = cell_right.add_paragraph()
        p2a.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2a.add_run("_" * 15).font.bold = True

        # Titles
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("\nĐỀ CƯƠNG CHI TIẾT HỌC PHẦN")
        run.font.bold = True
        run.font.size = Pt(14)

        # FIXED N-03: Bỏ dòng chú thích mẫu "(VIẾT TÊN...)", chỉ in tên thật
        p_name2 = self.doc.add_paragraph()
        p_name2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_name2 = p_name2.add_run(f"{(hp.get('ten_viet') or '').upper()}")
        run_name2.font.bold = True
        run_name2.font.size = Pt(14)

        p_level = self.doc.add_paragraph(f"Trình độ đào tạo: {hp.get('trinh_do') or 'Đại học'}")
        p_level.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_sec1_thong_tin_chung(self, hp):
        self._add_heading("1. Thông tin chung về học phần")
        self.doc.add_paragraph(f"Tên tiếng Việt: {(hp.get('ten_viet') or '').upper()}")
        self.doc.add_paragraph(f"Tên tiếng Anh: {(hp.get('ten_anh') or '').upper()}")
        self.doc.add_paragraph(f"Tên đơn vị quản lý học phần: {hp.get('don_vi_ql') or '...'}")
        self.doc.add_paragraph(f"Các giảng viên phụ trách học phần:")

        # FIXED M-01: Lấy đủ GV kể cả nhập tay (không có gv_id)
        # Dùng bảng hp_giang_vien trực tiếp, LEFT JOIN để không bỏ sót GV nhập tay
        def _get_gv_for_export(vai_tro):
            rows = self.db.conn.execute("""
                SELECT hpgv.ho_ten, hpgv.hoc_ham_vi, hpgv.sdt, hpgv.email,
                       gv.hoc_vi, gv.sdt AS gv_sdt, gv.email AS gv_email
                FROM hp_giang_vien hpgv
                LEFT JOIN giang_vien gv ON hpgv.gv_id = gv.id
                WHERE hpgv.hp_id=? AND hpgv.vai_tro=?
                ORDER BY hpgv.thu_tu
            """, (hp.get('id'), vai_tro)).fetchall()
            result = []
            for r in rows:
                result.append({
                    'ho_ten': r['ho_ten'] or '',
                    'hoc_vi': r['hoc_ham_vi'] or r['hoc_vi'] or '',
                    'sdt': r['sdt'] or r['gv_sdt'] or '',
                    'email': r['email'] or r['gv_email'] or ''
                })
            return result

        # 1.1 Table Giảng viên phụ trách chính
        self.doc.add_paragraph("Giảng viên phụ trách chính", style='Normal').runs[0].font.bold = True
        self._add_gv_table_from_list(_get_gv_for_export('phu_trach'))

        # 1.2 Table Giảng viên tham gia giảng dạy
        self.doc.add_paragraph("\nGiảng viên tham gia giảng dạy", style='Normal').runs[0].font.bold = True
        self._add_gv_table_from_list(_get_gv_for_export('tham_gia'))

        self.doc.add_paragraph("\n")

        # Info Table (Mã, Tín chỉ, Loại, Tính chất, Loại hình...)
        table = self.doc.add_table(rows=0, cols=2)
        table.width = Cm(16.5)
        
        info_rows = [
            ("Mã học phần: (mã học phần)", hp.get('ma') or hp.get('ma_hp')),
            ("Số tín chỉ: (Số TC)", hp.get('so_tin_chi')),
            ("Loại học phần:", hp.get('loai') or "Bắt buộc"),
            ("Tính chất học phần:", hp.get('tinh_chat') or "Lý thuyết"),
            ("Loại hình giảng dạy:", hp.get('loai_hinh') or "(Trực tiếp/Trực tuyến)"),
        ]
        
        for label, val in info_rows:
            row = table.add_row()
            row.cells[0].text = label
            row.cells[1].text = str(val or '')
            row.cells[0].paragraphs[0].runs[0].font.bold = True

        # FIXED M-02: Khởi tạo đúng số hàng (11) ngay từ đầu
        dist_data = [
            ("Giờ lên lớp", ""),
            (" + Lý thuyết, Bài tập, Kiểm tra", f"{hp.get('gio_lt') or 0} / {hp.get('gio_bt') or 0} / {hp.get('gio_kt') or 0}"),
            (" + Thực hành, Thí nghiệm", f"{hp.get('gio_th_tn') or 0}"),
            (" + Thảo luận (có nội dung)", f"{hp.get('gio_tl') or 0}"),
            (" + Tiểu luận, Đồ án", f"{hp.get('gio_tieu_luan') or 0}"),
            (" + Thực tập (tại doanh nghiệp, cssx, ..)", f"{hp.get('gio_thuc_tap') or 0}"),
            (" + Tự học, nghiên cứu, trải nghiệm", f"{hp.get('gio_tu_hoc') or 0}"),
            ("Tổng giờ học tập định mức", f"{hp.get('tong_gio') or 0}"),
            ("Học phần tiên quyết", hp.get('hp_tien_quyet') or "Không có"),
            ("Học phần song hành", hp.get('hp_song_hanh') or "Không có"),
            ("Học phần thay thế", hp.get('hp_thay_the') or "Không có")
        ]
        self.doc.add_paragraph("\nPhân bổ thời gian:")
        # Khởi tạo đúng số hàng ngay
        dist_table = self.doc.add_table(rows=len(dist_data), cols=2)
        dist_table.style = 'Table Grid'

        for i, (label, val) in enumerate(dist_data):
            row = dist_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = str(val or '')
            if not label.startswith(' '):  # Bold top-level headers
                if row.cells[0].paragraphs[0].runs:
                    row.cells[0].paragraphs[0].runs[0].font.bold = True

    def _add_gv_table(self, providers):
        """Legacy method — delegates to _add_gv_table_from_list."""
        converted = []
        for gv in providers:
            converted.append({
                'ho_ten': gv.get('ho_ten', ''),
                'hoc_vi': gv.get('hoc_vi', ''),
                'sdt': gv.get('sdt', ''),
                'email': gv.get('email', '')
            })
        self._add_gv_table_from_list(converted)

    def _add_gv_table_from_list(self, providers):
        """FIXED M-01: Render bảng GV từ list dict chuẩn hóa (hỗ trợ GV nhập tay)."""
        table = self.doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        for i, text in enumerate(["TT", "Học hàm, học vị, và tên", "Số điện thoại", "Email", "Vị trí"]):
            self._format_header_cell(hdr[i], text)

        if not providers:
            row = table.add_row()
            row.cells[0].text = "1"
            return

        for i, gv in enumerate(providers):
            row = table.add_row()
            row.cells[0].text = str(i + 1)
            hoc_vi = gv.get('hoc_vi', '') or ''
            ho_ten = gv.get('ho_ten', '') or ''
            row.cells[1].text = f"{hoc_vi} {ho_ten}".strip()
            row.cells[2].text = gv.get('sdt', '') or ''
            row.cells[3].text = gv.get('email', '') or ''
            row.cells[4].text = ''

    def _add_sec2_mo_ta(self, hp):
        self._add_heading("2. Mô tả tóm tắt nội dung học phần")
        self.doc.add_paragraph(hp.get('mo_ta') or "Chưa có mô tả.")

    def _add_sec3_muc_tieu(self, hp):
        self._add_heading("3. Mục tiêu học phần")
        self.doc.add_paragraph("(Mục tiêu học phần phải viết dưới góc độ người học và bắt đầu bằng một động từ hành động tương ứng with các cấp độ nắm vững kiến thức và có bổ ngữ làm rõ nghĩa cho động từ đó)")
        
        mts_rows = self.db.conn.execute("SELECT * FROM muc_tieu WHERE hp_id=? ORDER BY so_thu_tu", (hp.get('id'),)).fetchall()
        mts = [dict(r) for r in mts_rows]
        
        table = self.doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        for i, text in enumerate(["Mục tiêu\n(MT)", "Mô tả (mức độ tổng quát)", "CDR CTDT\n(PLO)"]):
            self._format_header_cell(hdr[i], text)
            
        for mt in mts:
            if mt.get('la_tieu_de_nhom'): continue # Skip group headers in table
            row = table.add_row()
            row.cells[0].text = f"MT{mt['so_thu_tu'] or ''}"
            row.cells[1].text = mt['mo_ta'] or ''
            row.cells[2].text = mt['cdr_ma'] or ''

    def _add_sec4_clo(self, hp):
        self._add_heading("4. Chuẩn đầu ra học phần (CLO)")
        self.doc.add_paragraph("(Tham khảo Quy định 975/QĐ-ĐHĐL... viết tường minh mục để đối chiếu với chuẩn đầu ra của CTĐT...)")
        
        clos_rows = self.db.conn.execute("SELECT * FROM clo WHERE hp_id=? ORDER BY id", (hp.get('id'),)).fetchall()
        clos = [dict(r) for r in clos_rows]
        
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        for i, text in enumerate(["CĐR học phần\n(CLO)", "Mô tả (mức độ chi tiết)", "CĐR CTDT\n(PLO)", "Mức độ\ngiảng dạy\n(I, R, M)"]):
            self._format_header_cell(hdr[i], text)
            
        for clo in clos:
            if clo.get('la_tieu_de_nhom'): continue
            row = table.add_row()
            row.cells[0].text = clo['ma'] or ''
            row.cells[1].text = clo['mo_ta'] or ''
            row.cells[2].text = clo['cdr_ma'] or ''
            row.cells[3].text = clo.get('level_irm') or ''

    def _add_sec5_hoc_lieu(self, hp):
        """FIXED N-09 -> 2026: Sec5 gồm 7 mục 5.1 đến 5.7: tài liệu + cơ sở vật chất"""
        self._add_heading("5. Cơ sở vật chất, trang thiết bị phục vụ dạy học")
        hls_rows = self.db.conn.execute("SELECT * FROM hoc_lieu WHERE hp_id=? ORDER BY loai, so_thu_tu", (hp.get('id'),)).fetchall()
        hls = [dict(r) for r in hls_rows]

        # 5.1
        p51 = self.doc.add_paragraph("5.1. Tài liệu học tập (Sách, giáo trình chính): ")
        if p51.runs: p51.runs[0].font.bold = True
        _tl_chinh = [h for h in hls if h['loai'] == '5.1']
        if _tl_chinh:
            for hl in _tl_chinh:
                self.doc.add_paragraph(hl['noi_dung'] or "", style='List Bullet')
        else:
            self.doc.add_paragraph("(Chưa có tài liệu chính)", style='List Bullet')

        # 5.2
        p52 = self.doc.add_paragraph("\n5.2. Tài liệu tham khảo:")
        if p52.runs: p52.runs[0].font.bold = True
        for hl in [h for h in hls if h['loai'] == '5.2']:
            self.doc.add_paragraph(hl['noi_dung'] or "", style='List Bullet')

        # 5.3
        p53 = self.doc.add_paragraph("\n5.3. Các tài liệu khác:")
        if p53.runs: p53.runs[0].font.bold = True
        for hl in [h for h in hls if h['loai'] == '5.3']:
            self.doc.add_paragraph(hl['noi_dung'] or "", style='List Bullet')
            
        # 5.4 Phòng học
        p54 = self.doc.add_paragraph("\n5.4. Phòng học:")
        if p54.runs: p54.runs[0].font.bold = True
        for hl in [h for h in hls if h['loai'] == '5.4']:
            self.doc.add_paragraph(hl['noi_dung'] or "", style='List Bullet')
            
        # 5.5 Thiết bị hỗ trợ
        p55 = self.doc.add_paragraph("\n5.5. Trang thiết bị hỗ trợ giảng dạy (nếu có):")
        if p55.runs: p55.runs[0].font.bold = True
        for hl in [h for h in hls if h['loai'] == '5.5']:
            self.doc.add_paragraph(hl['noi_dung'] or "", style='List Bullet')
            
        # 5.6 Thiết bị thực hành
        p56 = self.doc.add_paragraph("\n5.6. Thiết bị thực hành, thí nghiệm, PM hoặc phương tiện học tập khác (nếu có):")
        if p56.runs: p56.runs[0].font.bold = True
        for hl in [h for h in hls if h['loai'] == '5.6']:
            self.doc.add_paragraph(hl['noi_dung'] or "", style='List Bullet')
            
        # 5.7 Hoạt động ngoại khóa
        p57 = self.doc.add_paragraph("\n5.7. Các hoạt động ngoại khóa (nếu có):")
        if p57.runs: p57.runs[0].font.bold = True
        for hl in [h for h in hls if h['loai'] == '5.7']:
            self.doc.add_paragraph(hl['noi_dung'] or "", style='List Bullet')

    def _add_sec6_noi_dung(self, hp):
        self.is_phd = hp.get('nhom_hp_dac_thu') == 'Chuyên đề Tiến sĩ'
        self._add_heading("6. Nội dung chi tiết học phần")
        self.doc.add_paragraph("6.1. Phần lý thuyết")
        
        nd_lt_rows = self.db.conn.execute("SELECT * FROM noi_dung WHERE hp_id=? AND phan='lt' ORDER BY thu_tu", (hp.get('id'),)).fetchall()
        self._add_noi_dung_table([dict(r) for r in nd_lt_rows])

        self.doc.add_paragraph("\n6.2. Phần thực hành/thảo luận/đồ án/bài tập lớn")
        nd_th_rows = self.db.conn.execute("SELECT * FROM noi_dung WHERE hp_id=? AND phan='th' ORDER BY thu_tu", (hp.get('id'),)).fetchall()
        self._add_noi_dung_table([dict(r) for r in nd_th_rows])

    def _add_noi_dung_table(self, items):
        if not items:
            self.doc.add_paragraph("(Không có nội dung)")
            return

        table = self.doc.add_table(rows=2, cols=7)
        table.style = 'Table Grid'
        
        # Merge header for hours
        hdr0 = table.rows[0].cells
        hdr1 = table.rows[1].cells
        
        is_phd = getattr(self, 'is_phd', False)
        
        self._format_header_cell(hdr0[0], "TT"); hdr0[0].merge(hdr1[0])
        self._format_header_cell(hdr0[1], "Nội dung / Task" if is_phd else "Các nội dung cơ bản theo chương (bài), mục"); hdr0[1].merge(hdr1[1])
        self._format_header_cell(hdr0[2], "Nhiệm vụ cho NCS (*TS)" if is_phd else "Giờ lên lớp\n(LT/BT/\nTL/TH-TN)"); hdr0[2].merge(hdr1[2])
        self._format_header_cell(hdr0[3], "Hoạt động dạy và\nphương pháp"); hdr0[3].merge(hdr1[3])
        self._format_header_cell(hdr0[4], "Hoạt động học"); hdr0[4].merge(hdr1[4])
        self._format_header_cell(hdr0[5], "CĐR học phần\n(CLO)"); hdr0[5].merge(hdr1[5])
        self._format_header_cell(hdr0[6], "Bài đánh giá"); hdr0[6].merge(hdr1[6])

        idx = 1
        total_lt, total_bt, total_tl, total_th = 0, 0, 0, 0
        for nd in items:
            row = table.add_row()
            if nd['loai'] == 'chuong':
                row.cells[1].text = nd['ten'] or ''
                row.cells[1].paragraphs[0].runs[0].font.bold = True
            else:
                row.cells[0].text = str(idx)
                row.cells[1].text = nd['ten'] or ''
                if is_phd:
                    row.cells[2].text = nd.get('nhiem_vu_ncs', '') or ''
                else:
                    # formatted hours (LT/BT/TL/TH)
                    hrs = f"({nd['gio_lt'] or 0}/{nd['gio_bt'] or 0}/{nd['gio_tl'] or 0}/{nd['gio_th_tn'] or 0})"
                    row.cells[2].text = hrs
                row.cells[3].text = nd['pp_day'] or ''
                row.cells[4].text = nd['pp_hoc'] or ''
                row.cells[5].text = nd['cdr_ma'] or ''
                row.cells[6].text = nd['bai_danh_gia'] or ''
                
                total_lt += nd['gio_lt'] or 0
                total_bt += nd['gio_bt'] or 0
                total_tl += nd['gio_tl'] or 0
                total_th += nd['gio_th_tn'] or 0
                idx += 1
        
        # Add TỔNG row
        footer = table.add_row()
        footer.cells[1].text = "Tổng"
        footer.cells[1].paragraphs[0].runs[0].font.bold = True
        
        if is_phd:
            total_hrs = "(N/A)"
        else:
            total_hrs = f"({total_lt}/{total_bt}/{total_tl}/{total_th})"
        footer.cells[2].text = total_hrs
        footer.cells[2].paragraphs[0].runs[0].font.bold = True
        
        # Yellow background for the total row
        for cell in footer.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:fill'), 'FFFF00') # Yellow
            tcPr.append(shd)

    def _add_sec7_pp_day_hoc(self, hp):
        self._add_heading("7. Phương pháp giảng dạy và học tập")
        self.doc.add_paragraph(hp.get('pp_day_hoc') or "Theo phương pháp dạy học tích cực, kết hợp giữa lý thuyết và thực hành.")

    def _add_sec8_kiem_tra_danh_gia(self, hp):
        self._add_heading("8. Phương pháp, hình thức kiểm tra - đánh giá kết quả học tập học phần")
        
        # 8.1 Nhiệm vụ
        p81 = self.doc.add_paragraph()
        p81.add_run("8.1. Nhiệm vụ của sinh viên: ").font.bold = True
        self.doc.add_paragraph(hp.get('nhiem_vu_sv_len_lop') or "- Dự lớp đầy đủ...", style='List Bullet')
        
        # 8.2 Quy định
        p82 = self.doc.add_paragraph()
        p82.add_run("8.2. Quy định về thi cử, học vụ: ").font.bold = True
        self.doc.add_paragraph(hp.get('quy_dinh_hp') or "- Sinh viên phải dự lớp đầy đủ 70%...", style='List Bullet')

        # 8.3 Bảng đánh giá
        p83 = self.doc.add_paragraph()
        p83.add_run("8.3. Đánh giá kết quả học tập: ").font.bold = True
        
        bdgs_rows = self.db.conn.execute("SELECT * FROM ke_hoach_kiem_tra WHERE hp_id=? ORDER BY thu_tu", (hp.get('id'),)).fetchall()
        bdgs = [dict(r) for r in bdgs_rows]
        
        table = self.doc.add_table(rows=1, cols=8)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr_texts = ["Thành phần đánh giá", "Trọng số (%)", "Bài đánh giá", "Hình thức đánh giá", "Tiêu chí đánh giá", "CĐR học phần", "Điểm tối đa CĐR", "Trọng số CĐR (%)"]
        for i, text in enumerate(hdr_texts):
            self._format_header_cell(hdr[i], text)
            
        # Logical grouping and merging
        current_group = None
        group_start_row = None
        
        for b in bdgs:
            row = table.add_row()
            cells = row.cells
            
            group_name = b['nhom'] or "Khác"
            cells[0].text = group_name
            cells[1].text = f"{int((b['ty_trong_nhom'] or 0) * 100)}%"
            cells[2].text = b['noi_dung'] or ''
            cells[3].text = b['hinh_thuc'] or ''
            cells[4].text = b['tieu_chi_danh_gia'] or ''
            cells[5].text = b['clo_lien_quan'] or ''
            cells[6].text = b['diem_toi_da_cdr'] or '10'
            cells[7].text = f"{int((b['ty_trong'] or 0) * 100)}%"
            
            # Merging logic for the group name & group weight
            if group_name == current_group:
                # Merge with previous row(s)
                table.cell(group_start_row, 0).merge(cells[0])
                table.cell(group_start_row, 1).merge(cells[1])
            else:
                current_group = group_name
                group_start_row = table.rows.index(row)

        self._add_rubrics(hp)

    def _add_rubrics(self, hp):
        rubrics_rows = self.db.conn.execute("SELECT * FROM rubric_danh_gia WHERE hp_id=? ORDER BY thu_tu", (hp.get('id'),)).fetchall()
        rubrics = [dict(r) for r in rubrics_rows]
        if not rubrics: return

        self.doc.add_page_break()
        p = self.doc.add_paragraph()
        run = p.add_run("- Hệ thống rubrics:")
        run.font.bold = True
        run.font.size = Pt(13)

        for rb in rubrics:
            rb_title = self.doc.add_paragraph()
            rb_title.paragraph_format.space_before = Pt(12)
            run_rb = rb_title.add_run(f"RUBRIC {rb['ky_hieu'] or ''} – {rb['ten'] or ''}")
            run_rb.font.bold = True
            run_rb.font.color.rgb = RGBColor(255, 0, 0) # Red like in screenshot
            
            criteria_rows = self.db.conn.execute("SELECT * FROM rubric_tieu_chi WHERE rubric_id=? ORDER BY thu_tu", (rb['id'],)).fetchall()
            criteria = [dict(r) for r in criteria_rows]
            
            table = self.doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr_texts = ["Tiêu chí", "Trọng số", "Xuất sắc (9.0-10)", "Tốt (7.0-8.9)", "Đạt (5.0-6.9)", "Chưa đạt (0-4.9)"]
            for i, text in enumerate(hdr_texts):
                self._format_header_cell(hdr[i], text)
                if i > 1: # Red text for levels
                     hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)

            for cr in criteria:
                row = table.add_row()
                row.cells[0].text = cr['tieu_chi'] or ''
                row.cells[1].text = cr['trong_so'] or ''
                row.cells[2].text = cr['muc_xuat_sac'] or ''
                row.cells[3].text = cr['muc_tot'] or ''
                row.cells[4].text = cr['muc_dat'] or ''
                row.cells[5].text = cr['muc_chua_dat'] or ''
                
                # Format weight cell to be small and red if needed
                row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)

    def _add_sec9_quy_dinh(self, hp):
        """FIXED N-09: Sec9 chỉ xuất quy định học phần."""
        self._add_heading("9. Các quy định của học phần")
        self.doc.add_paragraph(
            hp.get('quy_dinh_hp') or "Tuân thủ các quy chế đào tạo hiện hành của Bộ GD&ĐT và Trường Đại học Điện lực."
        )

    def _add_sec10_co_so_phuc_vu(self, hp):
        # NOTE: 2026 mẫu không còn mục 10 này. Nó đã được nhập vào 5.
        pass

    def _add_sec11_doi_ngu_gv(self, hp):
        """FIXED M-01: Lấy đủ GV kể cả nhập tay."""
        self._add_heading("11. Đội ngũ giảng viên")
        self.doc.add_paragraph("Danh sách giảng viên tham gia giảng dạy và hướng dẫn học phần:")
        gv_rows = self.db.conn.execute("""
            SELECT hpgv.ho_ten, hpgv.hoc_ham_vi AS hoc_vi, hpgv.sdt, hpgv.email,
                   gv.hoc_vi AS gv_hoc_vi, gv.sdt AS gv_sdt, gv.email AS gv_email
            FROM hp_giang_vien hpgv
            LEFT JOIN giang_vien gv ON hpgv.gv_id = gv.id
            WHERE hpgv.hp_id=? ORDER BY hpgv.thu_tu
        """, (hp.get('id'),)).fetchall()
        gv_list = [{
            'ho_ten': r['ho_ten'] or '',
            'hoc_vi': r['hoc_vi'] or r['gv_hoc_vi'] or '',
            'sdt': r['sdt'] or r['gv_sdt'] or '',
            'email': r['email'] or r['gv_email'] or ''
        } for r in gv_rows]
        self._add_gv_table_from_list(gv_list)

    def _add_sec12_phu_luc(self, hp):
        self._add_heading("12. Phụ lục")
        self.doc.add_paragraph(hp.get('phu_luc') or "(Không có phụ lục)")

    def _add_sec13_cap_nhat(self, hp):
        self._add_heading("13. Lịch sử cập nhật")
        ls_rows = self.db.conn.execute("SELECT * FROM lich_su_cap_nhat WHERE hp_id=? ORDER BY lan DESC", (hp.get('id'),)).fetchall()
        if ls_rows:
            table = self.doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            for i, text in enumerate(["Lần", "Nội dung", "Quyết định", "Người CN", "Ngày CM"]):
                self._format_header_cell(hdr[i], text)
            for r in ls_rows:
                rd = dict(r)
                row = table.add_row()
                row.cells[0].text = str(rd['lan'])
                row.cells[1].text = rd['noi_dung'] or ''
                row.cells[2].text = rd['quyet_dinh'] or ''
                row.cells[3].text = rd['nguoi_cap_nhat'] or ''
                row.cells[4].text = rd['ngay_cap_nhat'] or ''
        else:
            self.doc.add_paragraph("(Chưa có lịch sử cập nhật)")

    def _add_sec14_chinh_sach(self, hp):
        self._add_heading("14. Chính sách học phần")
        cs_row = self.db.conn.execute("SELECT * FROM chinh_sach_hoc_phan WHERE hp_id=?", (hp.get('id'),)).fetchone()
        if not cs_row:
            self.doc.add_paragraph("(Không có dữ liệu chính sách)")
            return
        
        p1 = self.doc.add_paragraph()
        p1.add_run("- Liêm chính học thuật: ").font.bold = True
        p1.add_run(cs_row['liem_chinh_ht'] or '')
        
        p2 = self.doc.add_paragraph()
        p2.add_run("- Sử dụng AI: ").font.bold = True
        p2.add_run(cs_row['su_dung_ai'] or '')

    def _add_sec15_checklist(self, hp):
        self._add_heading("15. Checklist tự kiểm tra")
        cl_row = self.db.conn.execute("SELECT * FROM checklist_tu_kiem_tra WHERE hp_id=?", (hp.get('id'),)).fetchone()
        if not cl_row:
            self.doc.add_paragraph("(Không có checklist)")
            return
            
        items = [
            ("Hình thức (Font, Bảng, Dãn dòng chuẩn)", cl_row['hinh_thuc']),
            ("Chuẩn đầu ra Bloom (Không dùng Biết/Hiểu)", cl_row['clo_bloom']),
            ("Đảm bảo quy định 1 TC = 50h", cl_row['gio_tu_hoc']),
            ("100% CLO khớp nội dung kiểm tra", cl_row['rubric_match']),
            ("Giảng viên biên soạn đã ký xác nhận", cl_row['giang_vien_xn'])
        ]
        
        table = self.doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        for i, text in enumerate(["TT", "Hạng mục", "Trạng thái"]):
            self._format_header_cell(hdr[i], text)
            
        for i, (hang_muc, trang_thai) in enumerate(items):
            row = table.add_row()
            row.cells[0].text = str(i + 1)
            row.cells[1].text = hang_muc
            row.cells[2].text = 'Đạt' if trang_thai else 'Cần chú ý'

    def _add_signatures(self, hp):
        # Final Signatures
        self.doc.add_paragraph("\n")
        sig_table = self.doc.add_table(rows=1, cols=2)
        sig_left = sig_table.rows[0].cells[0]
        sig_right = sig_table.rows[0].cells[1]
        
        p_l = sig_left.add_paragraph("TRƯỞNG KHOA")
        p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_l.runs[0].font.bold = True
        sig_left.add_paragraph("\n\n\n")
        p_l_name = sig_left.add_paragraph(hp.get('ho_ten_ky_trai') or "(Họ và tên)")
        p_l_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_l_name.runs[0].font.bold = True
        
        p_r = sig_right.add_paragraph("NGƯỜI BIÊN SOẠN")
        p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_r.runs[0].font.bold = True
        sig_right.add_paragraph("\n\n\n")
        p_r_name = sig_right.add_paragraph(hp.get('ho_ten_ky_phai') or "(Họ và tên)")
        p_r_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_r_name.runs[0].font.bold = True

    def _add_heading(self, text):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.bold = True
        run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(12)

    def _format_header_cell(self, cell, text):
        cell.text = text
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.bold = True
        # Shading
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D9D9D9')
        tcPr.append(shd)

def export_de_cuong(db, hp_id_or_ma, output_path) -> bool:
    """Wrapper function to export course outline to Word."""
    try:
        exporter = DCCTHPExporter(db)
        exporter.export(hp_id_or_ma, output_path)
        return True
    except Exception as e:
        # Re-raise to allow controller to catch it, but print for console too
        print(f"Export error detail: {e}")
        raise e

# FIXED: Added missing functions expected by import_export_service.py
def export_builtin(db, hp_id, output_path):
    return export_de_cuong(db, hp_id, output_path)


def export_template(db, hp_id, out_path, template_path):
    """Xuất đề cương sử dụng jinja2 template (v2.0)."""
    try:
        from services.template_service import TemplateEngine
        
        engine = TemplateEngine(db)
        return engine.render(hp_id, template_path, out_path)
    except Exception as e:
        print(f"[Export] Template error: {e}")
        # Fallback to builtin if failed
        return export_de_cuong(db, hp_id, out_path)


def export_batch(db, hp_ids, dir_path, template_path=None, progress_callback=None):
    results = {'success': 0, 'errors': 0, 'details': []}
    total = len(hp_ids)
    
    for i, hid in enumerate(hp_ids):
        try:
            hp = db.get_hoc_phan(hid)
            name = hp['ten_viet'] if hp else f"HP_{hid}"
            ma = hp['ma'] if hp else str(hid)
            
            if progress_callback:
                progress_callback(i + 1, total, name, 'exporting')
            
            # Clean filename
            safe_name = "".join([c for c in f"{ma}_{name}" if c.isalnum() or c in (' ', '_', '-')]).strip()
            filename = f"{safe_name.replace(' ', '_')}.docx"
            out_path = os.path.join(dir_path, filename)
            
            if template_path:
                success = export_template(db, hid, out_path, template_path)
            else:
                success = export_builtin(db, hid, out_path)
                
            if success:
                results['success'] += 1
                results['details'].append({'hp_id': hid, 'status': 'ok', 'file': filename})
            else:
                results['errors'] += 1
                results['details'].append({'hp_id': hid, 'status': 'error', 'error': 'Export failed'})
                
        except Exception as e:
            results['errors'] += 1
            results['details'].append({'hp_id': hid, 'status': 'error', 'error': str(e)})
            
    return results
